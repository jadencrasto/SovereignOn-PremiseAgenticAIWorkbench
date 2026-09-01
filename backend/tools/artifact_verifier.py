"""
backend/tools/artifact_verifier.py
-----------------------------------
Post-generation verification tool for generated artifacts (DOCX, XLSX, CSV, Markdown, TXT).

Ensures deterministic correctness before concluding agent workflows:
1. Re-opens and parses generated file from disk.
2. For Word documents (.docx), validates genuine OOXML ZIP structure via python-docx.
3. Checks file integrity, non-zero size, and SHA-256 hash.
4. Validates required headers, schema consistency, and non-empty row/paragraph data.
5. Returns verified structured report to the agent memory and audit trail.
"""

from __future__ import annotations

import csv
import hashlib
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from pydantic import BaseModel, Field

from backend.tools.safety import validate_path_within

logger = logging.getLogger(__name__)


class ArtifactVerifierInput(BaseModel):
    """Input schema for the artifact_verifier tool."""
    relative_path: str = Field(
        ...,
        description="Path to the artifact to verify, e.g. 'data/sandbox/P-204_Maintenance_Summary.docx' or filename.",
    )
    expected_columns: Optional[List[str]] = Field(
        default=None,
        description="Optional list of column names or headings that must exist in the artifact.",
    )
    min_row_count: Optional[int] = Field(
        default=1,
        description="Minimum expected data rows or paragraphs.",
    )


def create_artifact_verifier(sandbox_dir: Path) -> callable:
    """Create the artifact verifier execution function."""

    async def execute_artifact_verifier(args: ArtifactVerifierInput) -> Dict[str, Any]:
        # Strip path prefixes if provided
        clean_name = args.relative_path.replace("data/sandbox/", "").replace("data\\sandbox\\", "").strip()
        target_path = validate_path_within(clean_name, sandbox_dir)

        if not target_path.exists():
            raise FileNotFoundError(f"Artifact not found on filesystem: {clean_name}")

        file_bytes = target_path.read_bytes()
        if len(file_bytes) == 0:
            raise ValueError(f"Artifact {clean_name} is empty (0 bytes).")

        sha256_hash = hashlib.sha256(file_bytes).hexdigest()
        suffix = target_path.suffix.lower()

        detected_headers: List[str] = []
        row_count = 0
        preview_rows: List[Any] = []
        doc_format = suffix.lstrip(".")
        extra_metadata: Dict[str, Any] = {}

        if suffix == ".docx":
            try:
                import docx

                doc = docx.Document(target_path)

                # Extract paragraphs
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                paragraph_count = len(paragraphs)
                table_count = len(doc.tables)

                for p in doc.paragraphs:
                    style_name = p.style.name.lower() if p.style else ""
                    if "heading" in style_name or "title" in style_name:
                        detected_headers.append(p.text.strip())

                # Inspect tables if present
                for tbl in doc.tables:
                    for r_idx, row in enumerate(tbl.rows):
                        row_vals = [c.text.strip() for c in row.cells]
                        if r_idx == 0 and not detected_headers:
                            detected_headers.extend(row_vals)
                        if len(preview_rows) < 5:
                            preview_rows.append(row_vals)

                # If no preview from tables, use paragraph preview
                if not preview_rows:
                    preview_rows = [[p] for p in paragraphs[:5]]

                row_count = max(paragraph_count, len(preview_rows), 1 if paragraph_count > 0 else 0)
                extra_metadata["paragraph_count"] = paragraph_count
                extra_metadata["table_count"] = table_count
                extra_metadata["preview_text"] = "\n".join(paragraphs[:3])

            except Exception as exc:
                raise ValueError(f"Corrupted or invalid DOCX document: {exc}")

        elif suffix == ".xlsx":
            try:
                import openpyxl
                wb = openpyxl.load_workbook(target_path, data_only=True)
                ws = wb.active

                # Find header row (row 4 in styled template or row 1)
                header_row_idx = 4 if ws.max_row >= 4 and ws.cell(row=4, column=1).value else 1
                for col in range(1, ws.max_column + 1):
                    val = ws.cell(row=header_row_idx, column=col).value
                    if val is not None:
                        detected_headers.append(str(val))

                # Count data rows
                for r in range(header_row_idx + 1, ws.max_row + 1):
                    row_vals = [ws.cell(row=r, column=c).value for c in range(1, len(detected_headers) + 1)]
                    if any(v is not None for v in row_vals):
                        row_count += 1
                        if len(preview_rows) < 5:
                            preview_rows.append(row_vals)

                wb.close()
            except Exception as exc:
                raise ValueError(f"Corrupted or invalid XLSX workbook: {exc}")

        elif suffix in (".csv", ".txt"):
            try:
                text_content = file_bytes.decode("utf-8")
                reader = csv.reader(text_content.splitlines())
                rows = list(reader)
                if rows:
                    detected_headers = rows[0]
                    data_rows = rows[1:]
                    row_count = len(data_rows)
                    preview_rows = data_rows[:5]
            except Exception as exc:
                raise ValueError(f"Failed to parse CSV artifact: {exc}")
        else:
            # Generic file
            row_count = 1
            detected_headers = ["file_content"]

        # Validate minimum row count
        if args.min_row_count is not None and row_count < args.min_row_count:
            raise ValueError(
                f"Artifact verification failed: found {row_count} rows/paragraphs, expected at least {args.min_row_count}."
            )

        # Validate expected columns
        missing_columns: List[str] = []
        if args.expected_columns:
            headers_lower = [h.lower() for h in detected_headers]
            for exp_col in args.expected_columns:
                if not any(exp_col.lower() in h for h in headers_lower):
                    missing_columns.append(exp_col)

            if missing_columns:
                raise ValueError(
                    f"Artifact verification failed: Missing required columns/headings: {missing_columns}. Found: {detected_headers}"
                )

        logger.info(
            "artifact_verified | path=%s rows=%d headers=%s hash=%s",
            clean_name, row_count, detected_headers, sha256_hash[:16]
        )

        res = {
            "verified": True,
            "filename": clean_name,
            "file_size_bytes": len(file_bytes),
            "sha256_hash": sha256_hash,
            "detected_headers": detected_headers,
            "row_count": row_count,
            "sample_preview": preview_rows,
            "format": doc_format,
            "status": "PASSED_VERIFICATION",
        }
        res.update(extra_metadata)
        return res

    return execute_artifact_verifier

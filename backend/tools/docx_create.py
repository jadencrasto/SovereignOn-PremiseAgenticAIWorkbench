"""
backend/tools/docx_create.py
------------------------------
Dedicated Microsoft Word (.docx) OOXML artifact generation tool.

Security & Integrity:
- Generates genuine ZIP-based OOXML .docx files using python-docx
- Writes strictly within data/sandbox/
- Atomic file generation via temporary files and os.replace()
- Path traversal and filename sanitization via backend.tools.safety
- Computes cryptographic SHA-256 hash of the generated artifact
- Mutating & high risk: requires human confirmation and planning approval
"""

from __future__ import annotations

import hashlib
import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from pydantic import BaseModel, Field

from backend.tools.safety import sanitize_filename, validate_path_within

logger = logging.getLogger(__name__)

_DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DocxCreateInput(BaseModel):
    """Input schema for creating genuine Word (.docx) documents."""
    filename: str = Field(
        ...,
        min_length=1,
        max_length=255,
        description="Target document filename ending in .docx, e.g. 'P-204_Maintenance_Summary.docx'.",
    )
    title: Optional[str] = Field(
        default="",
        description="Document title or top-level heading.",
    )
    content: Optional[str] = Field(
        default="",
        description="Document body text or markdown formatted content.",
    )
    paragraphs: Optional[List[str]] = Field(
        default=None,
        description="Optional list of paragraph strings to add to the document.",
    )
    sections: Optional[List[Dict[str, Any]]] = Field(
        default=None,
        description="Optional structured sections (e.g. [{'heading': '...', 'body': '...'}]).",
    )
    tables: Optional[List[Dict[str, Any]]] = Field(
        default=None,
        description="Optional tables to include (e.g. [{'headers': [...], 'rows': [[...]]}]).",
    )
    overwrite: bool = Field(
        default=True,
        description="Whether to overwrite existing file if present.",
    )


def create_docx_create(sandbox_dir: Path) -> callable:
    """
    Create the docx_create execution function bound to sandbox_dir.
    """
    sandbox_resolved = sandbox_dir.resolve()
    sandbox_resolved.mkdir(parents=True, exist_ok=True)

    async def execute_docx_create(args: DocxCreateInput) -> Dict[str, Any]:
        # 1. Sanitize and validate filename
        safe_name = sanitize_filename(args.filename)
        if not safe_name.lower().endswith(".docx"):
            safe_name += ".docx"

        target_path = validate_path_within(safe_name, sandbox_resolved)

        if target_path.exists() and not args.overwrite:
            raise ValueError(f"File '{safe_name}' already exists. Set overwrite=True to replace.")

        # 2. Build Document using python-docx
        try:
            import docx
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = docx.Document()

            # Add Title if present
            if args.title and args.title.strip():
                doc.add_heading(args.title.strip(), level=0)

            # Add Content (parse markdown-style lines)
            if args.content and args.content.strip():
                lines = args.content.strip().splitlines()
                current_p_lines: List[str] = []

                def flush_p():
                    if current_p_lines:
                        doc.add_paragraph(" ".join(current_p_lines).strip())
                        current_p_lines.clear()

                for line in lines:
                    trimmed = line.strip()
                    if not trimmed:
                        flush_p()
                        continue

                    if trimmed.startswith("### "):
                        flush_p()
                        doc.add_heading(trimmed[4:].strip(), level=3)
                    elif trimmed.startswith("## "):
                        flush_p()
                        doc.add_heading(trimmed[3:].strip(), level=2)
                    elif trimmed.startswith("# "):
                        flush_p()
                        doc.add_heading(trimmed[2:].strip(), level=1)
                    elif trimmed.startswith("- ") or trimmed.startswith("* "):
                        flush_p()
                        doc.add_paragraph(trimmed[2:].strip(), style="List Bullet")
                    else:
                        current_p_lines.append(trimmed)

                flush_p()

            # Add explicit paragraphs list if provided
            if args.paragraphs:
                for p_text in args.paragraphs:
                    if p_text and p_text.strip():
                        doc.add_paragraph(p_text.strip())

            # Add explicit structured sections if provided
            if args.sections:
                for sec in args.sections:
                    sec_heading = sec.get("heading")
                    sec_body = sec.get("body", "")
                    sec_level = sec.get("level", 1)
                    if sec_heading:
                        doc.add_heading(str(sec_heading), level=sec_level)
                    if sec_body:
                        doc.add_paragraph(str(sec_body))

            # Add explicit tables if provided
            if args.tables:
                for tbl_spec in args.tables:
                    headers = tbl_spec.get("headers", [])
                    rows = tbl_spec.get("rows", [])
                    if headers or rows:
                        num_cols = max(len(headers), max((len(r) for r in rows), default=0))
                        table = doc.add_table(rows=1 if headers else 0, cols=num_cols)
                        table.style = "Table Grid"

                        if headers:
                            hdr_cells = table.rows[0].cells
                            for i, h in enumerate(headers):
                                hdr_cells[i].text = str(h)

                        for r in rows:
                            row_cells = table.add_row().cells
                            for i, cell_val in enumerate(r):
                                if i < num_cols:
                                    row_cells[i].text = str(cell_val)

            # 3. Save to temporary file and replace atomically
            temp_path = target_path.with_suffix(".tmp.docx")
            doc.save(str(temp_path))

            os.replace(temp_path, target_path)

        except Exception as exc:
            logger.error("Docx generation failed: %s", exc)
            if "temp_path" in locals() and temp_path.exists():
                try:
                    temp_path.unlink()
                except OSError:
                    pass
            raise RuntimeError(f"Failed to generate DOCX document '{safe_name}': {exc}")

        # 4. Verify post-write integrity and compute SHA-256
        file_bytes = target_path.read_bytes()
        if len(file_bytes) == 0:
            raise RuntimeError(f"Generated DOCX document '{safe_name}' is empty on disk.")

        sha256_hash = hashlib.sha256(file_bytes).hexdigest()
        rel_path = f"data/sandbox/{safe_name}"

        logger.info(
            "docx_create | path=%s size=%d hash=%s",
            rel_path, len(file_bytes), sha256_hash[:16]
        )

        return {
            "created_path": rel_path,
            "filename": safe_name,
            "size_bytes": len(file_bytes),
            "sha256_hash": sha256_hash,
            "content_type": _DOCX_MIME_TYPE,
        }

    return execute_docx_create

"""
backend/rag/ingest.py
---------------------
Document parsing and chunking for the RAG pipeline.

Supported formats: PDF, TXT, MD, DOCX

Design:
  - Parsing produces a Document (full text + metadata).
  - Chunking splits a Document into Chunks with preserved metadata.
  - All chunk IDs are deterministic (doc_id + chunk index) so
    re-ingesting the same file produces the same IDs — safe for upsert.
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Internal data structures
# ---------------------------------------------------------------------------

@dataclass
class Document:
    """Parsed representation of an uploaded file."""
    document_id: str        # stable UUID derived from filename hash
    filename: str
    file_type: str          # pdf | txt | md | docx
    text: str               # full extracted text
    metadata: dict = field(default_factory=dict)


@dataclass
class Chunk:
    """A text chunk ready for embedding and storage."""
    chunk_id: str           # "{document_id}_chunk_{index}"
    document_id: str
    filename: str
    file_type: str
    text: str
    chunk_index: int
    metadata: dict = field(default_factory=dict)   # page, source, etc.


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

class DocumentParser:
    """
    Parses raw file bytes into a Document.

    Each parse_* method returns the extracted text and per-page metadata
    where available.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    @classmethod
    def parse(cls, filename: str, content: bytes) -> Document:
        """
        Main entry point.

        Args:
            filename : original uploaded filename (used for type detection)
            content  : raw file bytes

        Returns:
            Document with extracted text and metadata.

        Raises:
            ValueError : unsupported extension or empty content
        """
        suffix = Path(filename).suffix.lower()
        if suffix not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{suffix}'. "
                f"Supported: {sorted(cls.SUPPORTED_EXTENSIONS)}"
            )
        if not content:
            raise ValueError("File content is empty.")

        doc_id = cls._make_doc_id(filename, content)

        if suffix == ".pdf":
            text, metadata = cls._parse_pdf(content)
        elif suffix == ".docx":
            text, metadata = cls._parse_docx(content)
        else:
            # .txt and .md
            text, metadata = cls._parse_text(content)

        if not text.strip():
            raise ValueError(f"No text could be extracted from '{filename}'.")

        return Document(
            document_id=doc_id,
            filename=filename,
            file_type=suffix.lstrip("."),
            text=text,
            metadata=metadata,
        )

    # ------------------------------------------------------------------
    # Format-specific parsers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_pdf(content: bytes) -> tuple[str, dict]:
        """Extract text from PDF page by page."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("pypdf is required for PDF parsing: pip install pypdf")

        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"[Page {page_num}]\n{page_text.strip()}")

        text = "\n\n".join(pages)
        metadata = {"page_count": len(reader.pages), "source_format": "pdf"}
        logger.debug("PDF parsed: %d pages, %d chars", len(reader.pages), len(text))
        return text, metadata

    @staticmethod
    def _parse_docx(content: bytes) -> tuple[str, dict]:
        """Extract paragraphs from a DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise RuntimeError("python-docx is required: pip install python-docx")

        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        metadata = {
            "paragraph_count": len(paragraphs),
            "source_format": "docx",
        }
        # Extract core properties if available
        try:
            cp = doc.core_properties
            if cp.title:
                metadata["title"] = cp.title
            if cp.author:
                metadata["author"] = cp.author
        except Exception:
            pass
        logger.debug("DOCX parsed: %d paragraphs, %d chars", len(paragraphs), len(text))
        return text, metadata

    @staticmethod
    def _parse_text(content: bytes) -> tuple[str, dict]:
        """Read plain text / markdown, gracefully handling encoding."""
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                text = content.decode(encoding)
                return text, {"source_format": "text", "encoding": encoding}
            except UnicodeDecodeError:
                continue
        # Final fallback: replace errors
        text = content.decode("utf-8", errors="replace")
        logger.warning("Text file decoded with error replacement")
        return text, {"source_format": "text", "encoding": "utf-8-lossy"}

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _make_doc_id(filename: str, content: bytes) -> str:
        """Stable document ID from filename + content hash."""
        h = hashlib.sha256(filename.encode() + content[:1024]).hexdigest()[:16]
        return f"doc_{h}"


# ---------------------------------------------------------------------------
# Chunker
# ---------------------------------------------------------------------------

class TextChunker:
    """
    Splits a Document into overlapping text chunks.

    Strategy:
      1. Split text into paragraphs (double newline boundaries).
      2. Accumulate paragraphs into windows of approximately `chunk_size` chars.
      3. Apply `overlap` characters from the end of one chunk to the start
         of the next.

    This paragraph-aware approach preserves semantic boundaries better
    than pure character splitting.
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 150) -> None:
        if chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be less than chunk_size")
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk(self, document: Document) -> List[Chunk]:
        """
        Split a Document into Chunks.

        Returns an empty list only if the document had no usable text.
        """
        raw_paragraphs = self._split_paragraphs(document.text)
        windows = self._build_windows(raw_paragraphs)

        chunks: List[Chunk] = []
        for idx, window_text in enumerate(windows):
            text = window_text.strip()
            if not text:
                continue
            chunk_id = f"{document.document_id}_chunk_{idx}"
            # Carry page metadata if available (PDF stores "[Page N]" markers)
            page = self._extract_page_hint(text)
            chunk_meta = {**document.metadata, "chunk_index": idx}
            if page is not None:
                chunk_meta["page"] = page

            chunks.append(Chunk(
                chunk_id=chunk_id,
                document_id=document.document_id,
                filename=document.filename,
                file_type=document.file_type,
                text=text,
                chunk_index=idx,
                metadata=chunk_meta,
            ))

        logger.debug(
            "Chunked '%s': %d chunks (size=%d overlap=%d)",
            document.filename, len(chunks), self.chunk_size, self.chunk_overlap,
        )
        return chunks

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _split_paragraphs(text: str) -> List[str]:
        """Split on two or more newlines; keep non-empty paragraphs."""
        parts = re.split(r"\n{2,}", text)
        return [p.strip() for p in parts if p.strip()]

    def _build_windows(self, paragraphs: List[str]) -> List[str]:
        """
        Greedily accumulate paragraphs into windows ≤ chunk_size characters,
        then prepend an overlap tail from the previous window.
        """
        if not paragraphs:
            return []

        windows: List[str] = []
        current_parts: List[str] = []
        current_len = 0
        overlap_tail = ""

        for para in paragraphs:
            para_len = len(para)

            if current_len + para_len + 2 > self.chunk_size and current_parts:
                # Emit current window
                window = "\n\n".join(current_parts)
                if overlap_tail:
                    window = overlap_tail + "\n\n" + window
                windows.append(window)

                # Build overlap tail: tail characters of current window
                raw = "\n\n".join(current_parts)
                overlap_tail = raw[-self.chunk_overlap:] if len(raw) > self.chunk_overlap else raw

                current_parts = []
                current_len = 0

            current_parts.append(para)
            current_len += para_len + 2  # +2 for the separator

        # Flush the last window
        if current_parts:
            window = "\n\n".join(current_parts)
            if overlap_tail:
                window = overlap_tail + "\n\n" + window
            windows.append(window)

        return windows

    @staticmethod
    def _extract_page_hint(text: str) -> Optional[int]:
        """
        Extract the first [Page N] marker from text, if present.
        PDFs are parsed with these markers prepended per page.
        """
        m = re.search(r"\[Page (\d+)\]", text)
        if m:
            return int(m.group(1))
        return None

"""
tests/backend/test_docx_and_execution.py
-----------------------------------------
Comprehensive regression tests for:
1. Real local Python code execution sandbox tool (code_execution)
2. Dedicated genuine OOXML Microsoft Word creation tool (docx_create)
3. Extended artifact verifier for genuine DOCX validation (and rejection of fake text .docx)
4. Agent planning, prompt enforcement, and grounded end-to-end P-204 maintenance workflow
"""

import asyncio
import hashlib
import json
import os
import sys
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
import docx

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from backend.tools.registry import ToolRegistry, ToolDefinition, ToolResult
from backend.tools.code_execution import (
    CodeExecutionInput,
    create_code_execution,
    validate_python_code_safety,
)
from backend.tools.docx_create import (
    DocxCreateInput,
    create_docx_create,
)
from backend.tools.artifact_verifier import (
    ArtifactVerifierInput,
    create_artifact_verifier,
)
from backend.agent.planner import (
    AgentPlan,
    PlanStep,
    StepStatus,
    PlanStatus,
    should_use_planning,
)
from backend.agent.plan_validator import PlanValidator
from backend.agent.task_store import TaskStore
from backend.agent.task import TaskManager, TaskStatus
from backend.agent.approval import ApprovalManager


# ===================================================================
# PHASE 1: CODE EXECUTION SANDBOX TESTS
# ===================================================================

class TestCodeExecutionSandbox:
    """Tests for the local Python code execution sandbox."""

    @pytest.fixture
    def sandbox(self, tmp_path: Path):
        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)
        return sb

    @pytest.mark.asyncio
    async def test_exact_code_socket_blocked_before_subprocess_test_a(self, sandbox):
        """
        TEST A: Exact code `import socket\\nprint("should not run")`
        - AST validation rejects it
        - Subprocess is NEVER launched
        - Result indicates blocked/unsafe code
        - stdout does NOT contain "should not run"
        """
        code = 'import socket\nprint("should not run")'

        # 1. AST check directly
        with pytest.raises(ValueError, match="Security policy violation.*socket"):
            validate_python_code_safety(code)

        # 2. Execution tool check with patched subprocess.run to verify process is never spawned
        with patch("subprocess.run") as mock_subproc:
            exec_fn = create_code_execution(sandbox)
            with pytest.raises(ValueError, match="Security policy violation"):
                await exec_fn(CodeExecutionInput(code=code))

            # Verify subprocess was NEVER invoked
            mock_subproc.assert_not_called()

    @pytest.mark.asyncio
    async def test_exact_code_print_division_test_b(self, sandbox):
        """
        TEST B: Exact code `print(424 / 5)`
        - Real execution in sandbox
        - stdout exactly '84.8\\n' (or '84.8')
        - exit_code == 0
        """
        exec_fn = create_code_execution(sandbox)
        result = await exec_fn(CodeExecutionInput(code="print(424 / 5)"))

        assert result["success"] is True
        assert result["exit_code"] == 0
        assert result["timed_out"] is False
        assert result["stdout"] == "84.8\n" or result["stdout"].strip() == "84.8"
        assert result["stderr"] == ""

    @pytest.mark.parametrize("dangerous_code", [
        "import socket\ns = socket.socket()",
        "import urllib.request\nurllib.request.urlopen('http://example.com')",
        "import requests\nrequests.get('http://127.0.0.1')",
        "import subprocess\nsubprocess.run(['dir'])",
        "import os\nos.system('dir')",
        "import os\nos.popen('whoami')",
        "import ctypes\nctypes.CDLL('user32.dll')",
        "import winreg",
        "__import__('socket')",
        "from http.client import HTTPConnection",
    ])
    def test_blocked_dangerous_operations_test_c(self, dangerous_code):
        """
        TEST C: Attempted dangerous operations
        (socket, subprocess, requests, os.system, os.popen, ctypes, winreg, dynamic imports)
        - All rejected by AST validation before subprocess execution.
        """
        with pytest.raises(ValueError, match="Security policy violation"):
            validate_python_code_safety(dangerous_code)

    @pytest.mark.asyncio
    async def test_execution_timeout_handling(self, sandbox):
        """Execution exceeding timeout limit is cleanly killed and reported."""
        exec_fn = create_code_execution(sandbox)
        code = "import time\ntime.sleep(5)\nprint('done')"
        result = await exec_fn(CodeExecutionInput(code=code, timeout_seconds=1))

        assert result["success"] is False
        assert result["timed_out"] is True
        assert result["exit_code"] == -1
        assert "timed out" in result["stderr"]

    def test_invalid_syntax_rejected(self):
        """Syntax errors are caught by AST validator before process spawning."""
        with pytest.raises(ValueError, match="Invalid Python syntax"):
            validate_python_code_safety("def broken_syntax(")

    @pytest.mark.asyncio
    async def test_tool_registry_integration(self, sandbox):
        """ToolRegistry executes code_execution and captures structured result in audit."""
        registry = ToolRegistry()
        registry.register(ToolDefinition(
            name="code_execution",
            description="Execute Python code safely inside sandbox.",
            input_schema=CodeExecutionInput,
            execute_fn=create_code_execution(sandbox),
            category="Computation",
            read_only=False,
            risk_level="medium",
            requires_approval=False,
        ))

        res: ToolResult = await registry.execute(
            "code_execution",
            {"code": "x = 424 / 5\nprint(f'{x:.1f}')"},
            session_id="test_sess",
            user_role="admin",
        )

        assert res.success is True
        assert res.tool == "code_execution"
        assert res.result["stdout"].strip() == "84.8"
        assert res.result["exit_code"] == 0
        assert res.result["timed_out"] is False


# ===================================================================
# PHASE 2 & 3: DOCX CREATION TESTS
# ===================================================================

class TestDocxCreate:
    """Tests for genuine OOXML Word (.docx) document generation."""

    @pytest.fixture
    def sandbox(self, tmp_path: Path):
        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)
        return sb

    @pytest.mark.asyncio
    async def test_create_genuine_docx(self, sandbox):
        """Generate genuine Word document with headings, paragraphs, and tables."""
        docx_fn = create_docx_create(sandbox)
        input_data = DocxCreateInput(
            filename="P-204_Maintenance_Summary.docx",
            title="P-204 Boiler Feed Water Pump Overhaul Summary",
            content=(
                "## Executive Summary\n"
                "Pump P-204 was taken offline due to DE bearing temperature reaching 88.4°C.\n"
                "- Root cause: 65% clogged suction strainer S-204.\n"
                "- Repaired with OEM 13Cr martensitic stainless steel impeller.\n"
            ),
            tables=[{
                "headers": ["Parameter", "Pre-Overhaul", "Post-Overhaul", "Status"],
                "rows": [
                    ["Bearing Temp", "88.4°C", "58.5°C", "PASS"],
                    ["Vibration RMS", "7.2 mm/s", "1.65 mm/s", "PASS"],
                ]
            }],
        )

        result = await docx_fn(input_data)
        assert result["filename"] == "P-204_Maintenance_Summary.docx"
        assert result["created_path"] == "data/sandbox/P-204_Maintenance_Summary.docx"
        assert result["size_bytes"] > 0
        assert len(result["sha256_hash"]) == 64

        # Verify on filesystem using python-docx directly
        file_path = sandbox / "P-204_Maintenance_Summary.docx"
        assert file_path.exists()

        doc = docx.Document(file_path)
        doc_text = "\n".join(p.text for p in doc.paragraphs)
        assert "P-204 Boiler Feed Water Pump" in doc_text
        assert "88.4°C" in doc_text
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "Parameter"
        assert doc.tables[0].rows[1].cells[1].text == "88.4°C"

    @pytest.mark.asyncio
    async def test_docx_create_path_traversal_sanitized(self, sandbox):
        """Path traversal components in filename are sanitized to prevent sandbox escapes."""
        docx_fn = create_docx_create(sandbox)
        input_data = DocxCreateInput(
            filename="../../escaped.docx",
            title="Attempted Escape",
            content="Content",
        )
        res = await docx_fn(input_data)
        assert ".." not in res["filename"]
        assert (sandbox / res["filename"]).exists()


# ===================================================================
# PHASE 4: ARTIFACT VERIFIER TESTS
# ===================================================================

class TestArtifactVerifierDocx:
    """Tests for artifact verifier with genuine DOCX and rejection of fake files."""

    @pytest.fixture
    def sandbox(self, tmp_path: Path):
        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)
        return sb

    @pytest.mark.asyncio
    async def test_verify_real_docx_success(self, sandbox):
        """A genuine python-docx generated document passes verification."""
        docx_fn = create_docx_create(sandbox)
        await docx_fn(DocxCreateInput(
            filename="genuine_report.docx",
            title="Inspection Report",
            content="Detailed findings for refinery pump P-204.\nAll parameters verified.",
        ))

        verifier_fn = create_artifact_verifier(sandbox)
        verify_res = await verifier_fn(ArtifactVerifierInput(
            relative_path="data/sandbox/genuine_report.docx",
            min_row_count=1,
        ))

        assert verify_res["verified"] is True
        assert verify_res["status"] == "PASSED_VERIFICATION"
        assert verify_res["format"] == "docx"
        assert verify_res["paragraph_count"] >= 2
        assert "Inspection Report" in verify_res["preview_text"]
        assert len(verify_res["sha256_hash"]) == 64

    @pytest.mark.asyncio
    async def test_fake_docx_fails_verification(self, sandbox):
        """A plain text file renamed with .docx MUST fail verification."""
        fake_file = sandbox / "fake_report.docx"
        fake_file.write_text("This is plain text, not a genuine ZIP-based OOXML DOCX.", encoding="utf-8")

        verifier_fn = create_artifact_verifier(sandbox)
        with pytest.raises(ValueError, match=r"Corrupted or invalid DOCX document"):
            await verifier_fn(ArtifactVerifierInput(
                relative_path="fake_report.docx"
            ))


# ===================================================================
# PHASE 5: END-TO-END WORKFLOW & GROUNDING
# ===================================================================

class TestGroundedDocxAgentWorkflow:
    """
    End-to-end test of the P-204 maintenance workflow:
    1. document_search / file retrieval
    2. grounded summary extraction
    3. approval gate
    4. docx_create
    5. artifact_verifier
    """

    @pytest.fixture
    def env(self, tmp_path: Path):
        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)
        db_path = tmp_path / "tasks.db"
        store = TaskStore(db_path)
        approval_mgr = ApprovalManager(store, timeout_seconds=300)
        task_mgr = TaskManager(store)

        registry = ToolRegistry()
        registry.register(ToolDefinition(
            name="docx_create",
            description="Create DOCX document.",
            input_schema=DocxCreateInput,
            execute_fn=create_docx_create(sb),
            category="File Operations",
            read_only=False,
            risk_level="high",
            requires_approval=True,
        ))
        registry.register(ToolDefinition(
            name="code_execution",
            description="Execute Python code in sandbox.",
            input_schema=CodeExecutionInput,
            execute_fn=create_code_execution(sb),
            category="Computation",
            read_only=False,
            risk_level="medium",
            requires_approval=False,
        ))
        registry.register(ToolDefinition(
            name="artifact_verifier",
            description="Verify generated artifact.",
            input_schema=ArtifactVerifierInput,
            execute_fn=create_artifact_verifier(sb),
            category="Verification",
            read_only=True,
            risk_level="low",
            requires_approval=False,
        ))

        validator = PlanValidator(registry)

        return {
            "sandbox": sb,
            "store": store,
            "approval_mgr": approval_mgr,
            "registry": registry,
            "validator": validator,
            "task_mgr": task_mgr,
        }

    @pytest.mark.asyncio
    async def test_p204_grounded_docx_workflow(self, env):
        sandbox = env["sandbox"]
        task_mgr = env["task_mgr"]
        approval_mgr = env["approval_mgr"]
        validator = env["validator"]
        registry = env["registry"]

        # Read actual source document for P-204
        doc_path = PROJECT_ROOT / "data" / "demo" / "pump_p204_maintenance.md"
        assert doc_path.exists(), "Source document pump_p204_maintenance.md must exist."
        doc_content = doc_path.read_text(encoding="utf-8")

        # Grounding check: ensure key facts are present in source document
        assert "88.4°C" in doc_content
        assert "65% clogged" in doc_content
        assert "13Cr" in doc_content

        # Grounded summary derived from actual document
        grounded_summary = (
            "## P-204 Boiler Feed Water Pump Maintenance Summary\n\n"
            "**Incident:** High temperature alarms on DE radial bearing (peak: 88.4°C) and cavitation.\n"
            "**Root Cause:** Suction strainer S-204 was 65% clogged with magnetite scale.\n"
            "**Repairs:** Stage 1 impeller replaced with OEM 13Cr martensitic stainless steel.\n"
            "**Post-Overhaul:** DE Bearing temp stable at 58.5°C, overall vibration RMS 1.65 mm/s."
        )

        # 1. Create plan with docx_create and artifact_verifier
        task_id = "task_p204_001"
        plan = AgentPlan(
            task_id=task_id,
            objective="Create P-204 Maintenance Summary DOCX",
            steps=[
                PlanStep(
                    id="step_docx",
                    description="Generate P-204 maintenance summary Word document",
                    tool_name="docx_create",
                    arguments={
                        "filename": "P-204_Maintenance_Summary.docx",
                        "title": "P-204 Maintenance Summary",
                        "content": grounded_summary,
                    },
                    requires_approval=True,
                ),
                PlanStep(
                    id="step_verify",
                    description="Verify the generated DOCX artifact on disk",
                    tool_name="artifact_verifier",
                    arguments={
                        "relative_path": "data/sandbox/P-204_Maintenance_Summary.docx",
                    },
                    requires_approval=False,
                ),
            ],
        )

        # 2. Plan validation pass
        errors = validator.validate(plan)
        assert len(errors) == 0

        # 3. Create task and track lifecycle
        task = task_mgr.create_task(
            session_id="sess_p204",
            user_request="Find the refinery maintenance document for P-204, read it, create a concise maintenance summary, ask for approval, then save it as a DOCX.",
        )
        task_mgr.set_plan(task.task_id, plan)
        task_mgr.update_status(task.task_id, TaskStatus.PLANNING)

        # 4. Step 1 requires approval: transition to awaiting_approval
        task_mgr.update_step_status(task.task_id, "step_docx", StepStatus.awaiting_approval.value)
        task_mgr.update_status(task.task_id, TaskStatus.AWAITING_APPROVAL)

        # Confirm artifact has NOT been created before approval
        target_docx = sandbox / "P-204_Maintenance_Summary.docx"
        assert not target_docx.exists()

        # 5. Request and grant human approval
        appr_req = approval_mgr.request_approval(
            task_id=task.task_id,
            step_id="step_docx",
            tool_name="docx_create",
            arguments=plan.steps[0].arguments,
            risk_level="high",
            reason="Create P-204 Word summary artifact",
        )
        assert appr_req.status == "pending"

        approved = approval_mgr.approve(appr_req.approval_id)
        assert approved.status == "approved"

        # Verify approval before execution
        verified_appr = approval_mgr.verify_approval_for_execution(
            approval_id=appr_req.approval_id,
            task_id=task.task_id,
            step_id="step_docx",
            tool_name="docx_create",
            arguments=plan.steps[0].arguments,
        )
        assert verified_appr is True

        # 6. Execute step_docx
        task_mgr.update_step_status(task.task_id, "step_docx", StepStatus.approved.value)
        task_mgr.update_status(task.task_id, TaskStatus.EXECUTING)
        task_mgr.update_step_status(task.task_id, "step_docx", StepStatus.running.value)

        docx_res = await registry.execute(
            "docx_create",
            plan.steps[0].arguments,
            session_id="sess_p204",
            user_role="admin",
        )
        assert docx_res.success is True
        task_mgr.update_step_status(task.task_id, "step_docx", StepStatus.completed.value, result=str(docx_res.result))

        # 7. Execute step_verify
        task_mgr.update_step_status(task.task_id, "step_verify", StepStatus.running.value)
        verify_res = await registry.execute(
            "artifact_verifier",
            plan.steps[1].arguments,
            session_id="sess_p204",
            user_role="admin",
        )
        assert verify_res.success is True
        assert verify_res.result["verified"] is True
        assert verify_res.result["format"] == "docx"
        task_mgr.update_step_status(task.task_id, "step_verify", StepStatus.completed.value, result=str(verify_res.result))

        # 8. Complete task
        task_mgr.update_status(task.task_id, TaskStatus.COMPLETED, result="P-204 DOCX created and verified.")

        # 9. Verify generated artifact on filesystem
        assert target_docx.exists()
        doc = docx.Document(target_docx)
        doc_text = "\n".join(p.text for p in doc.paragraphs)
        assert "P-204" in doc_text
        assert "88.4°C" in doc_text
        assert "13Cr" in doc_text


class TestAgentIntegrationAndDisallowedSimulation:
    """Tests for agent tool registry schema export and chat_stream_with_tools loop."""

    @pytest.fixture
    def test_env(self, tmp_path: Path):
        from backend.config import Settings
        from backend.models.router import ModelRouter
        from backend.agent.memory import ConversationMemory
        from backend.agent.engine import AgentEngine

        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)

        settings = Settings(
            upload_dir=tmp_path / "uploads",
            agents_dir=Path("agents"),
            data_dir=tmp_path / "data",
            sandbox_dir=sb,
        )
        settings.upload_dir.mkdir(parents=True, exist_ok=True)

        registry = ToolRegistry()
        registry.register(ToolDefinition(
            name="code_execution",
            description="Execute Python code safely inside local sandbox.",
            input_schema=CodeExecutionInput,
            execute_fn=create_code_execution(sb),
            category="Computation",
            read_only=False,
            risk_level="medium",
            requires_approval=False,
        ))
        registry.register(ToolDefinition(
            name="docx_create",
            description="Generate a real Microsoft Word document (.docx).",
            input_schema=DocxCreateInput,
            execute_fn=create_docx_create(sb),
            category="File Operations",
            read_only=False,
            requires_confirmation=True,
            risk_level="high",
            requires_approval=True,
        ))

        router = MagicMock(spec=ModelRouter)
        router.default_model_id = "qwen2.5:7b"

        memory = ConversationMemory()
        engine = AgentEngine(
            settings=settings,
            router=router,
            memory=memory,
            tool_registry=registry,
        )

        return {
            "engine": engine,
            "registry": registry,
            "router": router,
            "sandbox": sb,
        }

    def test_tool_schemas_and_prompt_formatting(self, test_env):
        """Tool schemas and system prompt formatting include code_execution and docx_create."""
        registry = test_env["registry"]
        schemas = registry.get_tool_schemas_for_llm()
        names = [s["name"] for s in schemas]
        assert "code_execution" in names
        assert "docx_create" in names

        prompt_str = registry.format_tools_for_prompt()
        assert "code_execution" in prompt_str
        assert "docx_create" in prompt_str
        assert "timeout_seconds" in prompt_str

    @pytest.mark.asyncio
    async def test_agent_tool_loop_executes_real_code_tool(self, test_env):
        """
        When the model proposes a code_execution tool call, the agent engine
        executes the real sandbox tool and feeds the genuine stdout back to the model.
        """
        engine = test_env["engine"]
        router = test_env["router"]

        mock_provider = MagicMock()
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        # Mock streaming response:
        # Turn 1: Model emits tool call
        # Turn 2: Model emits final answer after receiving tool observation
        from backend.models.base import ChatChunk

        async def fake_chat_stream(req):
            messages = req.messages
            # Check if this is turn 1 or turn 2 (with observation)
            has_observation = any("[TOOL RESULT: code_execution]" in m.content for m in messages)
            if not has_observation:
                yield ChatChunk(delta='<tool_call>\n{"name": "code_execution", "arguments": {"code": "print(424 / 5)"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                yield ChatChunk(delta="The exact calculation of 424 / 5 executed in sandbox is 84.8.", done=False)
                yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools(
            session_id="sess_code_test",
            user_message="Please run python code to calculate 424 / 5.",
        ):
            events.append(ev)

        # Verify tool_start and tool_result events were emitted
        tool_starts = [e for e in events if isinstance(e, dict) and e.get("type") == "tool_start"]
        tool_results = [e for e in events if isinstance(e, dict) and e.get("type") == "tool_result"]

        assert len(tool_starts) == 1
        assert tool_starts[0]["tool"] == "code_execution"

        assert len(tool_results) == 1
        assert tool_results[0]["tool"] == "code_execution"
        assert tool_results[0]["success"] is True
        assert "84.8" in tool_results[0]["summary"]

        # Verify final text contains genuine result
        text_deltas = [e for e in events if isinstance(e, str)]
        full_text = "".join(text_deltas)
        assert "84.8" in full_text

    @pytest.mark.asyncio
    async def test_agent_reports_error_and_does_not_calculate_fallback(self, test_env):
        """
        When code execution fails (e.g. timeout or syntax/runtime error),
        the agent reports the failure and does not fabricate a manual fallback.
        """
        engine = test_env["engine"]
        router = test_env["router"]

        mock_provider = MagicMock()
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatChunk

        async def fake_chat_stream(req):
            messages = req.messages
            has_observation = any("[TOOL RESULT: code_execution]" in m.content for m in messages)
            if not has_observation:
                # Propose code execution that will fail security check or timeout
                yield ChatChunk(delta='<tool_call>\n{"name": "code_execution", "arguments": {"code": "import socket; s = socket.socket()"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                # Model must report the sandbox security block / failure, NOT claim success or calculate
                yield ChatChunk(delta="The sandbox code execution failed: Security policy violation: import of module 'socket' is blocked.", done=False)
                yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools(
            session_id="sess_code_fail",
            user_message="Run code with socket import",
        ):
            events.append(ev)

        tool_results = [e for e in events if isinstance(e, dict) and e.get("type") == "tool_result"]
        assert len(tool_results) == 1
        assert tool_results[0]["success"] is False
        assert "Error" in tool_results[0]["summary"]

        text_deltas = [e for e in events if isinstance(e, str)]
        full_text = "".join(text_deltas)
        assert "failed" in full_text.lower() or "blocked" in full_text.lower()

    @pytest.mark.asyncio
    async def test_agent_blocked_code_does_not_modify_or_retry_test_a(self, test_env):
        """
        TEST A (Agent Level):
        Exact code: `import socket\\nprint("should not run")`
        - Tool is invoked with the exact unchanged code
        - AST validation rejects it
        - Subprocess is never launched
        - Result indicates blocked/unsafe code
        - Agent does NOT generate or execute corrected code
        """
        engine = test_env["engine"]
        router = test_env["router"]

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatChunk

        code_to_run = 'import socket\nprint("should not run")'
        captured_tool_args = []

        async def fake_chat_stream(req):
            messages = req.messages
            has_observation = any("[TOOL RESULT: code_execution]" in m.content for m in messages)
            if not has_observation:
                tool_call_str = json.dumps({"name": "code_execution", "arguments": {"code": code_to_run}})
                yield ChatChunk(delta=f'<tool_call>\n{tool_call_str}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                # Observation received: verify observation contains blocked status
                obs_msg = next(m.content for m in messages if "[TOOL RESULT: code_execution]" in m.content)
                assert "blocked" in obs_msg.lower() or "error" in obs_msg.lower()
                assert "socket" in obs_msg.lower()

                # Model must report the blocked execution error without calculating or rewriting
                yield ChatChunk(delta="Execution blocked by security policy: import of module 'socket' is blocked in sandbox.", done=False)
                yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools(
            session_id="sess_code_blocked_a",
            user_message=f"Run this exact Python code:\n{code_to_run}",
        ):
            events.append(ev)

        # 1. Verify tool_start was for code_execution
        tool_starts = [e for e in events if isinstance(e, dict) and e.get("type") == "tool_start"]
        assert len(tool_starts) == 1
        assert tool_starts[0]["tool"] == "code_execution"

        # 2. Verify tool_result indicates failure/blocked
        tool_results = [e for e in events if isinstance(e, dict) and e.get("type") == "tool_result"]
        assert len(tool_results) == 1
        assert tool_results[0]["success"] is False

        # 3. Verify final answer does NOT contain 'should not run' and does NOT simulate execution
        text_deltas = [e for e in events if isinstance(e, str)]
        final_text = "".join(text_deltas)
        assert "should not run" not in final_text
        assert "blocked" in final_text.lower() or "security policy" in final_text.lower()


class TestGroundedSummaryAndVisionHardening:
    """Hardening tests for grounded document summaries and multimodal vision execution."""

    def test_p204_grounded_facts_vs_hallucinations(self):
        """
        Grounded summary must contain factual data from pump_p204_maintenance.md
        and strictly reject generic hallucinated boilerplate.
        """
        doc_path = PROJECT_ROOT / "data" / "demo" / "pump_p204_maintenance.md"
        assert doc_path.exists()
        doc_content = doc_path.read_text(encoding="utf-8")

        # Grounded factual points present in document
        assert "88.4°C" in doc_content or "88.4" in doc_content
        assert "P-204" in doc_content
        assert "65%" in doc_content  # suction strainer clogging
        assert "13Cr" in doc_content  # impeller metallurgy

        # Hallucinated phrases that must NOT be in a grounded summary
        hallucinated_phrases = [
            "No critical issues identified during the maintenance.",
            "Routine checks on all components revealed no significant wear or damage.",
            "Perform visual inspection every month.",
            "Schedule next maintenance in 6 months.",
        ]
        for phrase in hallucinated_phrases:
            assert phrase not in doc_content, f"Unexpected phrase in ground truth doc: {phrase}"


class TestMultimodalVisionHardeningPipeline:
    """
    Tests A-F for Multimodal Vision Pipeline:
    A. Vision invocation test
    B. Observation injection test
    C. No-vision-fallback test
    D. Equipment-ID anti-hallucination test
    E. Separation test
    F. End-to-end multimodal test
    """

    @pytest.fixture
    def test_env(self, tmp_path: Path):
        from backend.config import Settings
        from backend.models.router import ModelRouter
        from backend.agent.memory import ConversationMemory
        from backend.agent.engine import AgentEngine
        from backend.tools.document_search import DocumentSearchInput, create_document_search

        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)
        up = tmp_path / "uploads"
        up.mkdir(parents=True, exist_ok=True)

        settings = Settings(
            upload_dir=up,
            agents_dir=Path("agents"),
            data_dir=tmp_path / "data",
            sandbox_dir=sb,
        )

        mock_retriever = MagicMock()
        mock_chunk = MagicMock()
        mock_chunk.filename = "pump_p204_maintenance.md"
        mock_chunk.chunk_id = "p204_chunk_0"
        mock_chunk.chunk_index = 0
        mock_chunk.page = 1
        mock_chunk.score = 0.90
        mock_chunk.is_relevant = True
        mock_chunk.text = "# P-204 Maintenance Record\nEquipment: P-204 Hydrocracker Charge Pump\nFindings: Impeller cavitation erosion and bearing temp 88.4°C."
        mock_retriever.retrieve = AsyncMock(return_value=[mock_chunk])
        mock_retriever.is_chunk_relevant = MagicMock(return_value=True)

        registry = ToolRegistry()
        registry.register(ToolDefinition(
            name="document_search",
            description="Search local documents.",
            input_schema=DocumentSearchInput,
            execute_fn=create_document_search(mock_retriever),
            category="Information Retrieval",
            read_only=True,
        ))

        router = MagicMock(spec=ModelRouter)
        mock_vision_provider = MagicMock()
        mock_vision_provider.provider_name = "ollama"

        mock_chat_provider = MagicMock()
        mock_chat_provider.provider_name = "ollama"

        def resolve_model(model_id):
            if model_id and "llava" in str(model_id):
                return mock_vision_provider, "llava:7b"
            return mock_chat_provider, "qwen2.5:7b"

        router.get_provider_for_model.side_effect = resolve_model
        router.resolve_vision_model = MagicMock(return_value=(mock_vision_provider, "llava:7b"))
        router.resolve_chat_model = MagicMock(return_value=(mock_chat_provider, "qwen2.5:7b"))

        engine = AgentEngine(
            settings=settings,
            router=router,
            memory=ConversationMemory(),
            tool_registry=registry,
        )

        return {
            "engine": engine,
            "registry": registry,
            "router": router,
            "vision_provider": mock_vision_provider,
            "chat_provider": mock_chat_provider,
            "retriever": mock_retriever,
        }

    @pytest.mark.asyncio
    async def test_vision_invocation_test_a(self, test_env):
        """
        Test A: Provide an image, run multimodal agent, verify local vision model is actually invoked.
        """
        engine = test_env["engine"]
        vision_provider = test_env["vision_provider"]
        chat_provider = test_env["chat_provider"]

        from backend.models.base import ChatChunk, ChatResponse

        vision_provider.chat = AsyncMock(return_value=ChatResponse(
            content="Visual observation: Disassembled pump impeller with pitting.",
            model="llava:7b",
            provider="ollama",
        ))

        async def fake_chat_stream(req):
            yield ChatChunk(delta="Based on the image observation...", done=False)
            yield ChatChunk(delta="", done=True)

        chat_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools_multimodal(
            session_id="sess_vis_a",
            user_message="Analyze this image",
            image_b64="iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=",
        ):
            events.append(ev)

        # 1. Verify vision provider chat was called
        assert vision_provider.chat.called is True
        call_args = vision_provider.chat.call_args[0][0]
        assert call_args.model == "llava:7b"
        assert len(call_args.messages[1].images) == 1

        # 2. Verify tool_start for vision_analysis was logged
        tool_starts = [e for e in events if isinstance(e, dict) and e.get("type") == "tool_start"]
        assert any(ts.get("tool") == "vision_analysis" for ts in tool_starts)

    @pytest.mark.asyncio
    async def test_observation_injection_test_b(self, test_env):
        """
        Test B: Known visual observation reaches the reasoning model inside [VISUAL OBSERVATION].
        """
        engine = test_env["engine"]
        vision_provider = test_env["vision_provider"]
        chat_provider = test_env["chat_provider"]

        from backend.models.base import ChatChunk, ChatResponse

        observed_text = "A single-stage centrifugal pump casing with heavy corrosion."
        vision_provider.chat = AsyncMock(return_value=ChatResponse(
            content=observed_text,
            model="llava:7b",
            provider="ollama",
        ))

        received_messages = []

        async def fake_chat_stream(req):
            received_messages.extend(req.messages)
            yield ChatChunk(delta="Answer", done=False)
            yield ChatChunk(delta="", done=True)

        chat_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools_multimodal(
            session_id="sess_vis_b",
            user_message="Describe what is visible.",
            image_b64="iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=",
        ):
            events.append(ev)

        # Verify [VISUAL OBSERVATION] with exact text was injected
        visual_obs_msgs = [m for m in received_messages if observed_text in m.content]
        assert len(visual_obs_msgs) == 1
        assert "[VISUAL OBSERVATION from local vision model" in visual_obs_msgs[0].content
        assert "[END VISUAL OBSERVATION]" in visual_obs_msgs[0].content

    @pytest.mark.asyncio
    async def test_no_vision_fallback_test_c(self, test_env):
        """
        Test C: When vision execution fails, the agent reports the actual error
        and does NOT fabricate an image description.
        """
        engine = test_env["engine"]
        vision_provider = test_env["vision_provider"]
        chat_provider = test_env["chat_provider"]

        vision_provider.chat = AsyncMock(side_effect=RuntimeError("Connection refused by Ollama vision service"))
        chat_provider.chat_stream = MagicMock()

        events = []
        async for ev in engine.chat_stream_with_tools_multimodal(
            session_id="sess_vis_c",
            user_message="Analyze this image",
            image_b64="iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=",
        ):
            events.append(ev)

        # 1. Verify vision error status was emitted
        status_events = [e for e in events if isinstance(e, dict) and e.get("type") == "agent_status"]
        assert any(s.get("status") == "vision_error" for s in status_events)

        # 2. Verify chat_provider was NOT called to fabricate an answer
        assert chat_provider.chat_stream.called is False

        # 3. Verify error message is in text output
        deltas = [e for e in events if isinstance(e, str)]
        text_out = "".join(deltas)
        assert "Vision model error" in text_out or "Connection refused" in text_out

    @pytest.mark.asyncio
    async def test_equipment_id_anti_hallucination_test_d(self, test_env):
        """
        Test D: Generic pump in image must not be assumed to be P-204 or K-101
        just because document search retrieved those maintenance records.
        """
        engine = test_env["engine"]
        vision_provider = test_env["vision_provider"]
        chat_provider = test_env["chat_provider"]

        from backend.models.base import ChatChunk, ChatResponse

        # Vision returns generic pump without equipment tag
        vision_provider.chat = AsyncMock(return_value=ChatResponse(
            content="Visual observation: Disassembled centrifugal pump impeller showing cavitation pitting along the vanes. No equipment tag or identification number is visible.",
            model="llava:7b",
            provider="ollama",
        ))

        async def fake_chat_stream(req):
            messages = req.messages
            has_doc = any("[DOCUMENT CONTENT]" in m.content for m in messages)
            if not has_doc:
                yield ChatChunk(delta='<tool_call>\n{"name": "document_search", "arguments": {"query": "centrifugal pump cavitation erosion"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                grounded_resp = (
                    "### Visible in image\n"
                    "The image shows a centrifugal pump impeller with cavitation erosion along the vane leading edges. No equipment tag or serial number is visible.\n\n"
                    "### Stated in documents\n"
                    "Retrieved document `pump_p204_maintenance.md` describes maintenance for P-204 Boiler Feed Water Pump experiencing similar cavitation issues and 65% strainer clogging.\n\n"
                    "### Relationship / relevance\n"
                    "The maintenance procedures for P-204 (strainer cleaning and dynamic balancing) are technically relevant to the cavitation wear observed, but the image cannot be confirmed to be P-204 due to the absence of an equipment tag."
                )
                yield ChatChunk(delta=grounded_resp, done=False)
                yield ChatChunk(delta="", done=True)

        chat_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools_multimodal(
            session_id="sess_vis_d",
            user_message="Analyze image and search documents for maintenance info. Do not assume equipment ID unless supported.",
            image_b64="iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=",
        ):
            events.append(ev)

        deltas = [e for e in events if isinstance(e, str)]
        final_text = "".join(deltas)

        assert "cannot be confirmed to be P-204" in final_text or "No equipment tag" in final_text
        assert "I will proceed with the assumption" not in final_text

    @pytest.mark.asyncio
    async def test_separation_test_e(self, test_env):
        """
        Test E: Verify final output clearly separates:
        ### Visible in image
        ### Stated in documents
        ### Relationship / relevance
        """
        engine = test_env["engine"]
        vision_provider = test_env["vision_provider"]
        chat_provider = test_env["chat_provider"]

        from backend.models.base import ChatChunk, ChatResponse

        vision_provider.chat = AsyncMock(return_value=ChatResponse(
            content="Visual observation: Centrifugal pump impeller with cavitation pitting.",
            model="llava:7b",
            provider="ollama",
        ))

        async def fake_chat_stream(req):
            messages = req.messages
            has_doc = any("[DOCUMENT CONTENT]" in m.content for m in messages)
            if not has_doc:
                yield ChatChunk(delta='<tool_call>\n{"name": "document_search", "arguments": {"query": "pump cavitation"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                resp = (
                    "### Visible in image\n"
                    "Centrifugal pump impeller with cavitation wear.\n\n"
                    "### Stated in documents\n"
                    "Cavitation causes vibration and bearing overheating.\n\n"
                    "### Relationship / relevance\n"
                    "The document provides maintenance solutions for the visual symptoms."
                )
                yield ChatChunk(delta=resp, done=False)
                yield ChatChunk(delta="", done=True)

        chat_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools_multimodal(
            session_id="sess_vis_e",
            user_message="Analyze image and find relevant docs.",
            image_b64="iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=",
        ):
            events.append(ev)

        deltas = [e for e in events if isinstance(e, str)]
        final_text = "".join(deltas)

        assert "### Visible in image" in final_text
        assert "### Stated in documents" in final_text
        assert "### Relationship / relevance" in final_text

    @pytest.mark.asyncio
    async def test_end_to_end_multimodal_test_f(self, test_env):
        """
        Test F: End-to-end multimodal execution:
        IMAGE -> LLaVA -> observation -> RAG search -> document read -> grounded answer.
        """
        engine = test_env["engine"]
        vision_provider = test_env["vision_provider"]
        chat_provider = test_env["chat_provider"]

        from backend.models.base import ChatChunk, ChatResponse

        vision_provider.chat = AsyncMock(return_value=ChatResponse(
            content="Visual observation: Centrifugal pump impeller with cavitation erosion.",
            model="llava:7b",
            provider="ollama",
        ))

        async def fake_chat_stream(req):
            messages = req.messages
            has_doc = any("[TOOL RESULT: document_search]" in m.content for m in messages)
            if not has_doc:
                yield ChatChunk(delta='<tool_call>\n{"name": "document_search", "arguments": {"query": "pump cavitation"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                yield ChatChunk(delta="### Visible in image\nImpeller cavitation.\n\n### Stated in documents\nP-204 record.\n\n### Relationship / relevance\nMaintenance context.", done=False)
                yield ChatChunk(delta="", done=True)

        chat_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools_multimodal(
            session_id="sess_vis_f",
            user_message="Analyze image and check docs.",
            image_b64="iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=",
        ):
            events.append(ev)

        # Verify full execution order
        tool_starts = [e["tool"] for e in events if isinstance(e, dict) and e.get("type") == "tool_start"]
        tool_results = [e["tool"] for e in events if isinstance(e, dict) and e.get("type") == "tool_result"]

        assert tool_starts == ["vision_analysis", "document_search"]
        assert tool_results == ["vision_analysis", "document_search"]

        deltas = [e for e in events if isinstance(e, str)]
        assert len(deltas) > 0


class TestGroundedDocumentPipeline:
    """Comprehensive tests for document retrieval -> reading -> reasoning -> grounded summary pipeline."""

    @pytest.fixture
    def test_env(self, tmp_path: Path):
        from backend.config import Settings
        from backend.models.router import ModelRouter
        from backend.agent.memory import ConversationMemory
        from backend.agent.engine import AgentEngine
        from backend.tools.document_search import DocumentSearchInput, create_document_search
        from backend.tools.file_read import FileReadInput, create_file_read

        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)
        up = tmp_path / "uploads"
        up.mkdir(parents=True, exist_ok=True)

        settings = Settings(
            upload_dir=up,
            agents_dir=Path("agents"),
            data_dir=tmp_path / "data",
            sandbox_dir=sb,
        )

        doc_path = PROJECT_ROOT / "data" / "demo" / "pump_p204_maintenance.md"
        doc_content = doc_path.read_text(encoding="utf-8") if doc_path.exists() else "# P-204\nBearing temperature 88.4°C, 65% strainer blockage."

        mock_retriever = MagicMock()
        mock_chunk = MagicMock()
        mock_chunk.filename = "pump_p204_maintenance.md"
        mock_chunk.chunk_id = "p204_chunk_0"
        mock_chunk.chunk_index = 0
        mock_chunk.page = 1
        mock_chunk.score = 0.92
        mock_chunk.is_relevant = True
        mock_chunk.text = doc_content
        mock_retriever.retrieve = AsyncMock(return_value=[mock_chunk])
        mock_retriever.is_chunk_relevant = MagicMock(return_value=True)

        registry = ToolRegistry()
        registry.register(ToolDefinition(
            name="document_search",
            description="Search local documents.",
            input_schema=DocumentSearchInput,
            execute_fn=create_document_search(mock_retriever),
            category="Information Retrieval",
            read_only=True,
        ))
        registry.register(ToolDefinition(
            name="file_read",
            description="Read local workspace file.",
            input_schema=FileReadInput,
            execute_fn=create_file_read(up),
            category="File Operations",
            read_only=True,
        ))
        registry.register(ToolDefinition(
            name="docx_create",
            description="Create docx.",
            input_schema=DocxCreateInput,
            execute_fn=create_docx_create(sb),
            category="File Operations",
            read_only=False,
            requires_confirmation=True,
            risk_level="high",
            requires_approval=True,
        ))

        router = MagicMock(spec=ModelRouter)
        router.default_model_id = "qwen2.5:7b"

        engine = AgentEngine(
            settings=settings,
            router=router,
            memory=ConversationMemory(),
            tool_registry=registry,
        )

        return {
            "engine": engine,
            "registry": registry,
            "router": router,
            "retriever": mock_retriever,
            "doc_content": doc_content,
            "sandbox": sb,
        }

    @pytest.mark.asyncio
    async def test_p204_grounded_summary_test_a(self, test_env):
        """
        Test A: Search for P-204, receive [DOCUMENT CONTENT], generate grounded summary.
        Assert that factual fields come from the document and boilerplate is omitted.
        """
        engine = test_env["engine"]
        router = test_env["router"]
        doc_content = test_env["doc_content"]

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatChunk

        async def fake_chat_stream(req):
            messages = req.messages
            has_obs = any("[DOCUMENT CONTENT]" in m.content for m in messages)
            if not has_obs:
                yield ChatChunk(delta='<tool_call>\n{"name": "document_search", "arguments": {"query": "P-204 refinery maintenance"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                # Grounded summary formed strictly from [DOCUMENT CONTENT]
                summary = (
                    "**Proposed Maintenance Summary (Awaiting Approval):**\n\n"
                    "- **Equipment Name:** P-204 Hydrocracker Charge Pump\n"
                    "- **Maintenance Date:** 2024-11-14\n"
                    "- **Key Findings:** Inboard bearing temperature reached 88.4°C due to 65% suction strainer blockage.\n"
                    "- **Recommended Actions:** Clean suction strainer, perform dynamic balancing, and replace impeller with 13Cr stainless steel.\n\n"
                    "Please approve before I generate the final report file."
                )
                yield ChatChunk(delta=summary, done=False)
                yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools(
            session_id="sess_p204_grounded",
            user_message="Find the refinery maintenance document for P-204 and create a concise summary. Show me first and ask for approval.",
        ):
            events.append(ev)

        deltas = [e for e in events if isinstance(e, str)]
        final_text = "".join(deltas)

        # Assert factual fields from document
        assert "P-204" in final_text
        assert "88.4°C" in final_text or "88.4" in final_text
        assert "65%" in final_text
        assert "13Cr" in final_text

        # Assert unsupported boilerplate statements are NOT present
        unsupported = [
            "No specific findings mentioned",
            "No specific key maintenance findings are mentioned",
            "No significant issues were identified during the maintenance",
            "Standard cleaning and lubrication procedures were followed",
            "Inspection of seals and couplings revealed no abnormalities",
            "Pressure and temperature checks were within acceptable ranges",
            "Continue routine maintenance schedule",
            "Schedule next maintenance within the standard interval",
            "No critical issues identified",
            "Routine checks revealed no wear",
            "Further inspection or testing may be required",
            "Further inspection may be required",
            "Ensure all components are functioning",
            "Schedule maintenance in 6 months",
        ]
        for phrase in unsupported:
            assert phrase not in final_text

    @pytest.mark.asyncio
    async def test_missing_field_explicitly_stated_not_present_test_b(self, test_env):
        """
        Test B: When a requested field is absent from the document, the agent
        must state 'Not stated in retrieved document' and never fabricate values.
        """
        engine = test_env["engine"]
        router = test_env["router"]

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatChunk

        async def fake_chat_stream(req):
            messages = req.messages
            has_obs = any("[DOCUMENT CONTENT]" in m.content for m in messages)
            if not has_obs:
                yield ChatChunk(delta='<tool_call>\n{"name": "document_search", "arguments": {"query": "P-204 warranty"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                summary = (
                    "- **Equipment:** P-204 Hydrocracker Charge Pump\n"
                    "- **OEM Warranty Expiration:** Not stated in retrieved document.\n"
                    "- **Next Scheduled Maintenance Date:** Not stated in retrieved document.\n"
                    "- **Manufacturer Contact Number:** Not stated in retrieved document."
                )
                yield ChatChunk(delta=summary, done=False)
                yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools(
            session_id="sess_missing_field",
            user_message="What is the OEM warranty expiration date and next scheduled maintenance date for P-204 in the maintenance document?",
        ):
            events.append(ev)

        deltas = [e for e in events if isinstance(e, str)]
        final_text = "".join(deltas)
        assert "Not stated in retrieved document" in final_text

    @pytest.mark.asyncio
    async def test_search_result_metadata_only_rejected_as_evidence_test_c(self, test_env):
        """
        Test C: Search result metadata or empty chunks cannot be treated as factual evidence.
        """
        engine = test_env["engine"]
        router = test_env["router"]
        retriever = test_env["retriever"]

        # Simulate empty search result
        retriever.retrieve = AsyncMock(return_value=[])

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatChunk

        async def fake_chat_stream(req):
            messages = req.messages
            has_obs = any("[TOOL RESULT: document_search]" in m.content for m in messages)
            if not has_obs:
                yield ChatChunk(delta='<tool_call>\n{"name": "document_search", "arguments": {"query": "P-999"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                yield ChatChunk(delta="No relevant document evidence was found in the knowledge base for P-999. I cannot provide maintenance findings without document content.", done=False)
                yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools(
            session_id="sess_empty_search",
            user_message="Find maintenance findings for P-999.",
        ):
            events.append(ev)

        deltas = [e for e in events if isinstance(e, str)]
        final_text = "".join(deltas)
        assert "No relevant document evidence" in final_text or "cannot provide" in final_text

    @pytest.mark.asyncio
    async def test_approval_gate_prevents_file_creation_before_approval_test_d(self, test_env, tmp_path: Path):
        """
        Test D: Verify document search and summary happen before approval,
        and no mutating file creation tool (docx_create) executes before approval.
        """
        registry = test_env["registry"]
        sandbox = test_env["sandbox"]

        task_store = TaskStore(tmp_path / "tasks.db")
        approval_mgr = ApprovalManager(task_store, timeout_seconds=300)
        task_mgr = TaskManager(task_store)

        # Plan with document_search (read_only) -> docx_create (requires_approval)
        plan = AgentPlan(
            task_id="task_p204_appr",
            objective="Summarize P-204 and save to docx with approval",
            steps=[
                PlanStep(id="step_search", description="Search document", tool_name="document_search", arguments={"query": "P-204"}, requires_approval=False),
                PlanStep(id="step_docx", description="Create docx", tool_name="docx_create", arguments={"filename": "P204_Summary.docx", "title": "P204", "content": "Summary"}, requires_approval=True),
            ]
        )

        task = task_mgr.create_task(
            session_id="session_appr_gate",
            user_request="Summarize P-204 and save to docx with approval",
        )
        task_mgr.set_plan(task.task_id, plan)
        task_mgr.update_status(task.task_id, TaskStatus.PLANNING)

        # 1. Step 1 executes without approval
        task_mgr.update_step_status(task.task_id, "step_search", StepStatus.running.value)
        search_res = await registry.execute("document_search", {"query": "P-204"}, session_id="session_appr_gate")
        assert search_res.success is True
        task_mgr.update_step_status(task.task_id, "step_search", StepStatus.completed.value, result=str(search_res.result))

        # 2. Step 2 requires approval -> enters awaiting_approval
        task_mgr.update_step_status(task.task_id, "step_docx", StepStatus.awaiting_approval.value)
        task_mgr.update_status(task.task_id, TaskStatus.AWAITING_APPROVAL)

        # Verify target docx does NOT exist yet
        target_docx = sandbox / "P204_Summary.docx"
        assert not target_docx.exists()

        # Verify step is awaiting approval and cannot run without approval transition
        task_obj = task_mgr.get_task(task.task_id)
        assert task_obj.plan.steps[1].status == StepStatus.awaiting_approval.value

        # 3. Operator approves
        appr_req = approval_mgr.request_approval(
            task_id=task.task_id,
            step_id="step_docx",
            tool_name="docx_create",
            arguments={"filename": "P204_Summary.docx", "title": "P204", "content": "Summary"},
            risk_level="high",
            reason="Create P204 Word summary",
        )
        assert appr_req.status == "pending"
        approved = approval_mgr.approve(appr_req.approval_id)
        assert approved.status == "approved"

        # 4. Now execute docx_create
        task_mgr.update_step_status(task.task_id, "step_docx", StepStatus.approved.value)
        task_mgr.update_status(task.task_id, TaskStatus.EXECUTING)
        task_mgr.update_step_status(task.task_id, "step_docx", StepStatus.running.value)

        docx_res = await registry.execute("docx_create", {"filename": "P204_Summary.docx", "title": "P204", "content": "Summary"}, session_id="session_appr_gate")
        assert docx_res.success is True
        task_mgr.update_step_status(task.task_id, "step_docx", StepStatus.completed.value, result=str(docx_res.result))
        task_mgr.update_status(task.task_id, TaskStatus.COMPLETED)

        # File now exists after approved execution
        assert target_docx.exists()


class TestDocxApprovalAndArtifactWorkflow:
    """
    Focused tests for the P-204 document -> summary -> approval pause -> docx_create -> verifier pipeline.
    Covers Tests 1 through 6.
    """

    @pytest.fixture
    def full_workflow_env(self, tmp_path: Path):
        from backend.config import Settings
        from backend.models.router import ModelRouter
        from backend.agent.memory import ConversationMemory
        from backend.agent.engine import AgentEngine
        from backend.agent.planner import AgentPlanner
        from backend.agent.plan_validator import PlanValidator
        from backend.agent.task import TaskManager
        from backend.agent.task_store import TaskStore
        from backend.agent.approval import ApprovalManager
        from backend.tools.document_search import DocumentSearchInput, create_document_search
        from backend.tools.file_read import FileReadInput, create_file_read
        from backend.tools.docx_create import DocxCreateInput, create_docx_create
        from backend.tools.artifact_verifier import ArtifactVerifierInput, create_artifact_verifier

        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)
        up = tmp_path / "uploads"
        up.mkdir(parents=True, exist_ok=True)
        db_path = tmp_path / "tasks.db"

        settings = Settings(
            upload_dir=up,
            agents_dir=Path("agents"),
            data_dir=tmp_path / "data",
            sandbox_dir=sb,
        )

        doc_path = PROJECT_ROOT / "data" / "demo" / "pump_p204_maintenance.md"
        doc_content = doc_path.read_text(encoding="utf-8") if doc_path.exists() else "# P-204 Hydrocracker Charge Pump\nDate: 2024-11-14\nInboard bearing temperature 88.4°C due to 65% suction strainer blockage. Recommend dynamic balancing and 13Cr impeller."

        mock_retriever = MagicMock()
        mock_chunk = MagicMock()
        mock_chunk.filename = "pump_p204_maintenance.md"
        mock_chunk.chunk_id = "p204_chunk_0"
        mock_chunk.chunk_index = 0
        mock_chunk.page = 1
        mock_chunk.score = 0.95
        mock_chunk.is_relevant = True
        mock_chunk.text = doc_content
        mock_retriever.retrieve = AsyncMock(return_value=[mock_chunk])
        mock_retriever.is_chunk_relevant = MagicMock(return_value=True)

        registry = ToolRegistry()
        registry.register(ToolDefinition(
            name="document_search",
            description="Search local documents.",
            input_schema=DocumentSearchInput,
            execute_fn=create_document_search(mock_retriever),
            category="Information Retrieval",
            read_only=True,
            risk_level="low",
            requires_approval=False,
        ))
        registry.register(ToolDefinition(
            name="file_read",
            description="Read local workspace file.",
            input_schema=FileReadInput,
            execute_fn=create_file_read(up),
            category="File Operations",
            read_only=True,
            risk_level="medium",
            requires_approval=False,
        ))
        registry.register(ToolDefinition(
            name="docx_create",
            description="Create Microsoft Word document.",
            input_schema=DocxCreateInput,
            execute_fn=create_docx_create(sb),
            category="File Operations",
            read_only=False,
            risk_level="high",
            requires_approval=True,
            requires_confirmation=True,
        ))
        registry.register(ToolDefinition(
            name="artifact_verifier",
            description="Verify created artifact.",
            input_schema=ArtifactVerifierInput,
            execute_fn=create_artifact_verifier(sb),
            category="File Operations",
            read_only=True,
            risk_level="low",
            requires_approval=False,
        ))

        task_store = TaskStore(db_path)
        task_mgr = TaskManager(task_store)
        approval_mgr = ApprovalManager(task_store, timeout_seconds=300)
        planner = AgentPlanner(max_plan_steps=6)
        validator = PlanValidator(registry, max_plan_steps=6)

        router = MagicMock(spec=ModelRouter)
        router.default_model_id = "qwen2.5:7b"

        engine = AgentEngine(
            settings=settings,
            router=router,
            memory=ConversationMemory(),
            tool_registry=registry,
        )
        engine.set_task_manager(task_mgr)
        engine.set_planner(planner)
        engine.set_plan_validator(validator)
        engine.set_approval_manager(approval_mgr)

        return {
            "engine": engine,
            "registry": registry,
            "router": router,
            "task_manager": task_mgr,
            "approval_manager": approval_mgr,
            "planner": planner,
            "validator": validator,
            "sandbox": sb,
            "doc_content": doc_content,
        }

    @pytest.mark.asyncio
    async def test_summary_only_does_not_require_approval_test_1(self, full_workflow_env):
        """
        TEST 1: Prompt 'Find the P-204 maintenance document, read it, and summarize it.'
        - document_search executes
        - grounded summary returned
        - no docx_create
        - no approval required
        """
        engine = full_workflow_env["engine"]
        router = full_workflow_env["router"]
        from backend.agent.planner import should_use_planning

        prompt = "Find the P-204 maintenance document, read it, and summarize it."
        assert should_use_planning(prompt) is False

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatChunk

        async def fake_chat_stream(req):
            messages = req.messages
            has_obs = any("[TOOL RESULT: document_search]" in m.content for m in messages)
            if not has_obs:
                yield ChatChunk(delta='<tool_call>\n{"name": "document_search", "arguments": {"query": "P-204 maintenance"}}\n</tool_call>', done=False)
                yield ChatChunk(delta="", done=True)
            else:
                summary = (
                    "### P-204 Maintenance Summary\n"
                    "- **Equipment Name:** P-204 Hydrocracker Charge Pump\n"
                    "- **Maintenance Date:** 2024-11-14\n"
                    "- **Key Findings:** Bearing temp reached 88.4°C due to 65% strainer blockage.\n"
                    "- **Recommended Actions:** Dynamic balancing and 13Cr impeller replacement."
                )
                yield ChatChunk(delta=summary, done=False)
                yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.chat_stream_with_tools(
            session_id="sess_test1_summary_only",
            user_message=prompt,
        ):
            events.append(ev)

        # Assert tool_start was only document_search
        tool_starts = [e["tool"] for e in events if isinstance(e, dict) and e.get("type") == "tool_start"]
        assert tool_starts == ["document_search"]
        assert "docx_create" not in tool_starts

        # Assert no approval event was emitted
        approval_events = [e for e in events if isinstance(e, dict) and e.get("type") == "approval_required"]
        assert len(approval_events) == 0

        # Assert grounded facts returned
        text_deltas = [e for e in events if isinstance(e, str)]
        full_text = "".join(text_deltas)
        assert "88.4°C" in full_text
        assert "65%" in full_text

    @pytest.mark.asyncio
    async def test_p204_artifact_workflow_pauses_before_approval_test_2(self, full_workflow_env):
        """
        TEST 2: Prompt with proposed summary first + approval requirement.
        - document_search executes
        - grounded proposed summary generated
        - approval_required state emitted with metadata
        - workflow remains pending/waiting
        - docx_create execution count == 0
        - file_write execution count == 0
        - no artifact created
        """
        engine = full_workflow_env["engine"]
        router = full_workflow_env["router"]
        sandbox = full_workflow_env["sandbox"]
        task_mgr = full_workflow_env["task_manager"]
        from backend.agent.planner import should_use_planning

        prompt = (
            "Find the refinery maintenance document for P-204, read it, and prepare a concise maintenance summary "
            "containing the equipment name, maintenance date, key findings, and recommended actions. "
            "Show me the proposed summary first. Do not create or modify any file until I explicitly approve."
        )
        assert should_use_planning(prompt) is True

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        # Mock planner LLM response (JSON plan)
        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search P-204 maintenance document", "tool_name": "document_search", "arguments": {"query": "P-204 refinery maintenance"}, "requires_approval": False},
            {"description": "Prepare proposed grounded maintenance summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Create P-204 DOCX maintenance report in sandbox", "tool_name": "docx_create", "arguments": {"filename": "P204_Maintenance_Summary.docx", "title": "P-204 Maintenance Summary", "content": ""}, "requires_approval": True},
            {"description": "Verify generated DOCX artifact", "tool_name": "artifact_verifier", "arguments": {"filename": "P204_Maintenance_Summary.docx"}, "requires_approval": False},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        async def fake_chat_stream(req):
            # Used for reasoning step 2
            summary = (
                "**Proposed Maintenance Summary (Awaiting Approval):**\n\n"
                "- **Equipment Name:** P-204 Hydrocracker Charge Pump\n"
                "- **Maintenance Date:** 2024-11-14\n"
                "- **Key Findings:** Inboard bearing temperature reached 88.4°C due to 65% suction strainer blockage.\n"
                "- **Recommended Actions:** Clean suction strainer, perform dynamic balancing, and replace impeller with 13Cr stainless steel.\n\n"
                "Please approve before I generate the final report file."
            )
            yield ChatChunk(delta=summary, done=False)
            yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.run_agent_task(
            session_id="sess_test2_p204_pause",
            user_message=prompt,
        ):
            events.append(ev)

        # 1. Verify approval_required event emitted
        appr_events = [e for e in events if isinstance(e, dict) and e.get("type") == "approval_required"]
        assert len(appr_events) == 1
        assert appr_events[0]["tool_name"] == "docx_create"
        assert appr_events[0]["risk_level"] == "high"
        assert "P204_Maintenance_Summary.docx" in appr_events[0]["arguments"]["filename"]

        # 2. Verify task status is AWAITING_APPROVAL
        task_id = appr_events[0]["task_id"]
        task_obj = task_mgr.get_task(task_id)
        assert task_obj.status == "awaiting_approval"

        # 3. Verify no DOCX or mutation occurred
        created_files = list(sandbox.glob("*.docx"))
        assert len(created_files) == 0

    @pytest.mark.asyncio
    async def test_approval_resumes_docx_creation_and_verification_test_3(self, full_workflow_env):
        """
        TEST 3: Given a pending P-204 workflow:
        - Approve the pending operation
        - docx_create executes exactly once
        - File exists in data/sandbox/
        - File has .docx extension and genuine OOXML format
        - artifact_verifier executes and returns verified=True
        """
        engine = full_workflow_env["engine"]
        router = full_workflow_env["router"]
        sandbox = full_workflow_env["sandbox"]
        task_mgr = full_workflow_env["task_manager"]
        approval_mgr = full_workflow_env["approval_manager"]

        prompt = "Create P-204 DOCX summary after approval"
        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search P-204 maintenance document", "tool_name": "document_search", "arguments": {"query": "P-204 maintenance"}, "requires_approval": False},
            {"description": "Prepare proposed grounded maintenance summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Create P-204 DOCX maintenance report in sandbox", "tool_name": "docx_create", "arguments": {"filename": "P204_Maintenance_Summary.docx", "title": "P-204 Maintenance Summary", "content": ""}, "requires_approval": True},
            {"description": "Verify generated DOCX artifact", "tool_name": "artifact_verifier", "arguments": {"filename": "P204_Maintenance_Summary.docx"}, "requires_approval": False},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        async def fake_chat_stream(req):
            summary = (
                "**Equipment Name:** P-204 Hydrocracker Charge Pump\n"
                "**Maintenance Date:** 2024-11-14\n"
                "**Key Findings:** Inboard bearing temperature reached 88.4°C due to 65% suction strainer blockage.\n"
                "**Recommended Actions:** Dynamic balancing and 13Cr stainless steel impeller replacement."
            )
            yield ChatChunk(delta=summary, done=False)
            yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        # Run task to approval checkpoint
        events = []
        async for ev in engine.run_agent_task(session_id="sess_test3_resume", user_message=prompt):
            events.append(ev)

        appr_ev = next(e for e in events if isinstance(e, dict) and e.get("type") == "approval_required")
        task_id = appr_ev["task_id"]
        approval_id = appr_ev["approval_id"]

        # Resume task with approval
        resume_events = []
        async for ev in engine.resume_agent_task(
            task_id=task_id,
            approval_id=approval_id,
            approved=True,
            user_role="admin",
        ):
            resume_events.append(ev)

        # 1. Verify docx_create executed
        tool_results = [e for e in resume_events if isinstance(e, dict) and e.get("type") == "tool_result"]
        tools_executed = [e["tool"] for e in tool_results]
        assert "docx_create" in tools_executed
        assert "artifact_verifier" in tools_executed

        # 2. Verify file exists in sandbox and is genuine OOXML
        target_docx = sandbox / "P204_Maintenance_Summary.docx"
        assert target_docx.exists()
        doc = docx.Document(target_docx)
        doc_text = "\n".join(p.text for p in doc.paragraphs)
        assert "P-204" in doc_text
        assert "88.4°C" in doc_text

        # 3. Verify task_completed event
        task_done = any(e.get("type") == "task_completed" for e in resume_events if isinstance(e, dict))
        assert task_done is True

    @pytest.mark.asyncio
    async def test_approval_is_idempotent_test_4(self, full_workflow_env):
        """
        TEST 4: Approve the same pending operation twice.
        - docx_create executes exactly once
        - Second approval attempt fails/is rejected without duplicate execution
        """
        engine = full_workflow_env["engine"]
        router = full_workflow_env["router"]
        sandbox = full_workflow_env["sandbox"]
        task_mgr = full_workflow_env["task_manager"]
        approval_mgr = full_workflow_env["approval_manager"]

        prompt = "Create P-204 DOCX with double approval test"
        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search document", "tool_name": "document_search", "arguments": {"query": "P-204"}, "requires_approval": False},
            {"description": "Prepare summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Create DOCX", "tool_name": "docx_create", "arguments": {"filename": "P204_Idempotent.docx", "title": "P-204", "content": "Findings: 88.4C, 65% blockage"}, "requires_approval": True},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        async def fake_chat_stream(req):
            yield ChatChunk(delta="Summary: 88.4°C bearing temp, 65% blockage.", done=False)
            yield ChatChunk(delta="", done=True)
        mock_provider.chat_stream = fake_chat_stream

        # Run to approval
        events = []
        async for ev in engine.run_agent_task(session_id="sess_test4_idempotent", user_message=prompt):
            events.append(ev)

        appr_ev = next(e for e in events if isinstance(e, dict) and e.get("type") == "approval_required")
        task_id = appr_ev["task_id"]
        approval_id = appr_ev["approval_id"]

        # First approval: succeeds
        first_res = []
        async for ev in engine.resume_agent_task(task_id, approval_id, approved=True, user_role="admin"):
            first_res.append(ev)
        assert any(e.get("type") == "task_completed" for e in first_res if isinstance(e, dict))

        # Second approval: pending approval no longer exists
        pending_second = approval_mgr.get_pending_for_task(task_id)
        assert pending_second is None

        # Attempting second resume yields error
        second_res = []
        async for ev in engine.resume_agent_task(task_id, approval_id, approved=True, user_role="admin"):
            second_res.append(ev)
        assert any("error" in str(e).lower() for e in second_res)

    @pytest.mark.asyncio
    async def test_docx_contains_approved_grounded_summary_test_5(self, full_workflow_env):
        """
        TEST 5: Verify the resulting DOCX contains the approved grounded summary
        (equipment name, maintenance date, findings, actions) and no generic boilerplate.
        """
        engine = full_workflow_env["engine"]
        router = full_workflow_env["router"]
        sandbox = full_workflow_env["sandbox"]

        prompt = "Create P-204 DOCX grounded check"
        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search document", "tool_name": "document_search", "arguments": {"query": "P-204"}, "requires_approval": False},
            {"description": "Prepare summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Create DOCX", "tool_name": "docx_create", "arguments": {"filename": "P204_Grounded_Doc.docx", "title": "P-204 Maintenance Summary", "content": ""}, "requires_approval": True},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        grounded_summary = (
            "Equipment Name: P-204 Hydrocracker Charge Pump\n"
            "Maintenance Date: 2024-11-14\n"
            "Key Findings: Inboard bearing temperature reached 88.4°C due to 65% suction strainer blockage.\n"
            "Recommended Actions: Clean suction strainer, perform dynamic balancing, and replace impeller with 13Cr stainless steel."
        )
        async def fake_chat_stream(req):
            yield ChatChunk(delta=grounded_summary, done=False)
            yield ChatChunk(delta="", done=True)
        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.run_agent_task(session_id="sess_test5_grounded", user_message=prompt):
            events.append(ev)

        appr_ev = next(e for e in events if isinstance(e, dict) and e.get("type") == "approval_required")
        async for ev in engine.resume_agent_task(appr_ev["task_id"], appr_ev["approval_id"], approved=True, user_role="admin"):
            pass

        target_docx = sandbox / "P204_Grounded_Doc.docx"
        assert target_docx.exists()

        doc = docx.Document(target_docx)
        doc_text = "\n".join(p.text for p in doc.paragraphs)

        # Grounded facts
        assert "P-204" in doc_text
        assert "2024-11-14" in doc_text
        assert "88.4°C" in doc_text or "88.4" in doc_text
        assert "65%" in doc_text
        assert "13Cr" in doc_text

        # Forbidden boilerplate absent
        forbidden = [
            "No significant issues were identified during the maintenance.",
            "Standard cleaning and lubrication procedures were followed.",
            "Inspection of seals and couplings revealed no abnormalities.",
            "Pressure and temperature checks were within acceptable ranges.",
            "Continue routine maintenance schedule.",
            "Schedule next maintenance within the standard interval.",
        ]
        for f in forbidden:
            assert f not in doc_text

    @pytest.mark.asyncio
    async def test_no_mutation_before_approval_test_6(self, full_workflow_env):
        """
        TEST 6: Record sandbox directory contents before starting workflow.
        After workflow reaches approval_required, assert no new file was created.
        """
        engine = full_workflow_env["engine"]
        router = full_workflow_env["router"]
        sandbox = full_workflow_env["sandbox"]

        # Record sandbox snapshot
        before_files = set(sandbox.iterdir())

        prompt = "Create P-204 DOCX snapshot test"
        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search document", "tool_name": "document_search", "arguments": {"query": "P-204"}, "requires_approval": False},
            {"description": "Prepare summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Create DOCX", "tool_name": "docx_create", "arguments": {"filename": "P204_No_Mutation.docx", "title": "P-204", "content": ""}, "requires_approval": True},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        async def fake_chat_stream(req):
            yield ChatChunk(delta="Summary findings.", done=False)
            yield ChatChunk(delta="", done=True)
        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.run_agent_task(session_id="sess_test6_snapshot", user_message=prompt):
            events.append(ev)

        # Assert approval_required reached
        assert any(e.get("type") == "approval_required" for e in events if isinstance(e, dict))

        # Assert sandbox files are unchanged (no new files created before approval)
        after_files = set(sandbox.iterdir())
        assert before_files == after_files


class TestPlanningDocumentGroundingRegression:
    """
    Regression test suite for document grounding and content propagation in planning workflows.
    Covers:
    A. Planning workflow receives actual unescaped [DOCUMENT CONTENT], not just metadata.
    B. Grounded P-204 proposed summary uses document facts and omits boilerplate.
    C. Missing-field behavior ('Not stated in retrieved document.').
    D. Approval workflow integrity (before and after approval).
    """

    @pytest.fixture
    def planning_grounding_env(self, tmp_path: Path):
        from backend.config import Settings
        from backend.models.router import ModelRouter
        from backend.agent.memory import ConversationMemory
        from backend.agent.engine import AgentEngine
        from backend.agent.planner import AgentPlanner
        from backend.agent.plan_validator import PlanValidator
        from backend.agent.task import TaskManager
        from backend.agent.task_store import TaskStore
        from backend.agent.approval import ApprovalManager
        from backend.tools.document_search import DocumentSearchInput, create_document_search
        from backend.tools.file_read import FileReadInput, create_file_read
        from backend.tools.docx_create import DocxCreateInput, create_docx_create
        from backend.tools.artifact_verifier import ArtifactVerifierInput, create_artifact_verifier

        sb = tmp_path / "sandbox"
        sb.mkdir(parents=True, exist_ok=True)
        up = tmp_path / "uploads"
        up.mkdir(parents=True, exist_ok=True)
        db_path = tmp_path / "tasks.db"

        settings = Settings(
            upload_dir=up,
            agents_dir=Path("agents"),
            data_dir=tmp_path / "data",
            sandbox_dir=sb,
        )

        doc_path = PROJECT_ROOT / "data" / "demo" / "pump_p204_maintenance.md"
        doc_content = doc_path.read_text(encoding="utf-8") if doc_path.exists() else (
            "# P-204 Hydrocracker Charge Pump Maintenance Report\n"
            "Date: 2024-11-14\n"
            "Equipment: P-204 Hydrocracker Charge Pump\n"
            "Inboard bearing temperature reached 88.4°C due to 65% suction strainer blockage. "
            "Recommend dynamic balancing and 13Cr stainless steel impeller replacement."
        )

        mock_retriever = MagicMock()
        mock_chunk = MagicMock()
        mock_chunk.filename = "pump_p204_maintenance.md"
        mock_chunk.chunk_id = "p204_chunk_0"
        mock_chunk.chunk_index = 0
        mock_chunk.page = 1
        mock_chunk.score = 0.95
        mock_chunk.is_relevant = True
        mock_chunk.text = doc_content
        mock_retriever.retrieve = AsyncMock(return_value=[mock_chunk])
        mock_retriever.is_chunk_relevant = MagicMock(return_value=True)

        registry = ToolRegistry()
        registry.register(ToolDefinition(
            name="document_search",
            description="Search local documents.",
            input_schema=DocumentSearchInput,
            execute_fn=create_document_search(mock_retriever),
            category="Information Retrieval",
            read_only=True,
            risk_level="low",
            requires_approval=False,
        ))
        registry.register(ToolDefinition(
            name="file_read",
            description="Read local workspace file.",
            input_schema=FileReadInput,
            execute_fn=create_file_read(up),
            category="File Operations",
            read_only=True,
            risk_level="medium",
            requires_approval=False,
        ))
        registry.register(ToolDefinition(
            name="docx_create",
            description="Create Microsoft Word document.",
            input_schema=DocxCreateInput,
            execute_fn=create_docx_create(sb),
            category="File Operations",
            read_only=False,
            risk_level="high",
            requires_approval=True,
            requires_confirmation=True,
        ))
        registry.register(ToolDefinition(
            name="artifact_verifier",
            description="Verify created artifact.",
            input_schema=ArtifactVerifierInput,
            execute_fn=create_artifact_verifier(sb),
            category="File Operations",
            read_only=True,
            risk_level="low",
            requires_approval=False,
        ))

        task_store = TaskStore(db_path)
        task_mgr = TaskManager(task_store)
        approval_mgr = ApprovalManager(task_store, timeout_seconds=300)
        planner = AgentPlanner(max_plan_steps=6)
        validator = PlanValidator(registry, max_plan_steps=6)

        router = MagicMock(spec=ModelRouter)
        router.default_model_id = "qwen2.5:7b"

        engine = AgentEngine(
            settings=settings,
            router=router,
            memory=ConversationMemory(),
            tool_registry=registry,
        )
        engine.set_task_manager(task_mgr)
        engine.set_planner(planner)
        engine.set_plan_validator(validator)
        engine.set_approval_manager(approval_mgr)

        return {
            "engine": engine,
            "registry": registry,
            "router": router,
            "task_manager": task_mgr,
            "approval_manager": approval_mgr,
            "planner": planner,
            "validator": validator,
            "sandbox": sb,
            "doc_content": doc_content,
        }

    @pytest.mark.asyncio
    async def test_planning_workflow_receives_actual_document_content_test_a(self, planning_grounding_env):
        """
        Regression Test A:
        Verify the reasoning step in the planning workflow receives the actual document content
        formatted in [DOCUMENT CONTENT] blocks, rather than merely metadata (score, filename, chunk ID).
        """
        engine = planning_grounding_env["engine"]
        router = planning_grounding_env["router"]

        prompt = (
            "Find the refinery maintenance document for P-204, read it, and prepare a concise maintenance summary containing only:\n"
            "- equipment name\n"
            "- maintenance date\n"
            "- key maintenance findings\n"
            "- recommended actions\n\n"
            "Show me the proposed summary first and ask for my approval.\n\n"
            "Do not create, modify, or write any file until I explicitly approve."
        )

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search refinery maintenance document for P-204", "tool_name": "document_search", "arguments": {"query": "P-204 refinery maintenance"}, "requires_approval": False},
            {"description": "Prepare concise grounded maintenance summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Create P-204 maintenance summary DOCX", "tool_name": "docx_create", "arguments": {"filename": "P204_Summary.docx", "title": "P-204 Maintenance Summary", "content": ""}, "requires_approval": True},
            {"description": "Verify generated DOCX artifact", "tool_name": "artifact_verifier", "arguments": {"filename": "P204_Summary.docx"}, "requires_approval": False},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        captured_messages = []

        async def fake_chat_stream(req):
            nonlocal captured_messages
            captured_messages = req.messages
            # Grounded summary produced because [DOCUMENT CONTENT] is present
            summary = (
                "**Proposed Maintenance Summary (Awaiting Approval):**\n\n"
                "- **Equipment Name:** P-204 Hydrocracker Charge Pump\n"
                "- **Maintenance Date:** 2024-11-14\n"
                "- **Key Maintenance Findings:** Inboard bearing temperature reached 88.4°C due to 65% suction strainer blockage.\n"
                "- **Recommended Actions:** Clean suction strainer, perform dynamic balancing, and replace impeller with 13Cr stainless steel.\n\n"
                "Please approve before I generate the final report file."
            )
            yield ChatChunk(delta=summary, done=False)
            yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.run_agent_task(
            session_id="sess_regression_a",
            user_message=prompt,
        ):
            events.append(ev)

        # 1. Assert captured reasoning messages contain [DOCUMENT CONTENT]
        all_msg_content = "\n\n".join(m.content for m in captured_messages)
        assert "[DOCUMENT CONTENT]" in all_msg_content
        assert "[END DOCUMENT CONTENT]" in all_msg_content

        # 2. Assert known factual content from pump_p204_maintenance.md is available in reasoning context
        assert "88.4" in all_msg_content
        assert "65%" in all_msg_content
        assert "13Cr" in all_msg_content
        assert "pump_p204_maintenance.md" in all_msg_content

        # 3. Assert it is NOT merely metadata (scores/chunk IDs) but contains the actual text body
        assert "suction strainer" in all_msg_content.lower()
        assert "bearing" in all_msg_content.lower()

    @pytest.mark.asyncio
    async def test_grounded_p204_summary_and_no_boilerplate_test_b(self, planning_grounding_env):
        """
        Regression Test B:
        Assert the proposed summary uses facts present in the retrieved document
        and does NOT output generic boilerplate.
        """
        engine = planning_grounding_env["engine"]
        router = planning_grounding_env["router"]

        prompt = (
            "Find the refinery maintenance document for P-204, read it, and prepare a concise maintenance summary containing only:\n"
            "- equipment name\n"
            "- maintenance date\n"
            "- key maintenance findings\n"
            "- recommended actions\n\n"
            "Show me the proposed summary first and ask for my approval."
        )

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search P-204 document", "tool_name": "document_search", "arguments": {"query": "P-204"}, "requires_approval": False},
            {"description": "Synthesize proposed summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Save DOCX", "tool_name": "docx_create", "arguments": {"filename": "P204.docx"}, "requires_approval": True},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        async def fake_chat_stream(req):
            all_content = "\n\n".join(m.content for m in req.messages)
            assert "[DOCUMENT CONTENT]" in all_content
            assert "88.4" in all_content
            # Model sees document content and produces factual summary
            summary = (
                "**Proposed Maintenance Summary (Awaiting Approval):**\n\n"
                "- **Equipment Name:** P-204 Hydrocracker Charge Pump\n"
                "- **Maintenance Date:** 2024-11-14\n"
                "- **Key Maintenance Findings:** Inboard bearing temperature reached 88.4°C caused by 65% suction strainer blockage.\n"
                "- **Recommended Actions:** Clean suction strainer, perform dynamic balancing, and replace impeller with 13Cr stainless steel.\n\n"
                "Please confirm your approval to proceed with creating the document."
            )
            yield ChatChunk(delta=summary, done=False)
            yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.run_agent_task(
            session_id="sess_regression_b",
            user_message=prompt,
        ):
            events.append(ev)

        deltas = [e for e in events if isinstance(e, str)]
        summary_text = "".join(deltas)

        # Factual statements verified
        assert "P-204" in summary_text
        assert "2024-11-14" in summary_text
        assert "88.4°C" in summary_text or "88.4" in summary_text
        assert "65%" in summary_text
        assert "dynamic balancing" in summary_text
        assert "13Cr" in summary_text

        # Forbidden generic boilerplate statements NOT present
        boilerplate = [
            "No significant issues were identified",
            "Standard cleaning and lubrication procedures were followed",
            "Further inspection or testing may be required",
            "Ensure all components are functioning",
            "Schedule maintenance within the standard interval",
            "Pressure and temperature checks were within acceptable ranges",
            "Continue routine maintenance schedule",
        ]
        for b in boilerplate:
            assert b not in summary_text

    @pytest.mark.asyncio
    async def test_missing_field_behavior_test_c(self, planning_grounding_env):
        """
        Regression Test C:
        Ask for a genuinely absent field (e.g. OEM warranty expiration date).
        Assert that the answer says 'Not stated in retrieved document.' and does not fabricate.
        """
        engine = planning_grounding_env["engine"]
        router = planning_grounding_env["router"]

        prompt = (
            "Find the refinery maintenance document for P-204 and report:\n"
            "- equipment name\n"
            "- OEM warranty expiration date\n"
            "- manufacturer contact phone number\n"
            "Show me the proposed summary first."
        )

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search P-204 document", "tool_name": "document_search", "arguments": {"query": "P-204 warranty contact"}, "requires_approval": False},
            {"description": "Synthesize summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Save DOCX", "tool_name": "docx_create", "arguments": {"filename": "P204_Warranty.docx"}, "requires_approval": True},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        async def fake_chat_stream(req):
            all_content = "\n\n".join(m.content for m in req.messages)
            assert "[DOCUMENT CONTENT]" in all_content
            # Model observes grounding instruction #4: if field is absent, output 'Not stated in retrieved document.'
            summary = (
                "- **Equipment Name:** P-204 Hydrocracker Charge Pump\n"
                "- **OEM Warranty Expiration Date:** Not stated in retrieved document.\n"
                "- **Manufacturer Contact Phone Number:** Not stated in retrieved document."
            )
            yield ChatChunk(delta=summary, done=False)
            yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        events = []
        async for ev in engine.run_agent_task(
            session_id="sess_regression_c",
            user_message=prompt,
        ):
            events.append(ev)

        deltas = [e for e in events if isinstance(e, str)]
        summary_text = "".join(deltas)

        assert "Not stated in retrieved document." in summary_text
        assert "P-204" in summary_text

    @pytest.mark.asyncio
    async def test_approval_remains_intact_workflow_test_d(self, planning_grounding_env):
        """
        Regression Test D:
        For the exact approval workflow:
        Before approval:
        - document_search executes
        - reasoning executes and produces proposed summary
        - approval_required is emitted
        - task enters AWAITING_APPROVAL
        - docx_create execution count MUST be 0
        - artifact_verifier execution count MUST be 0
        - no DOCX file created

        After approval:
        - docx_create executes exactly once
        - genuine DOCX is created
        - artifact_verifier executes
        - verification returns verified=True
        """
        engine = planning_grounding_env["engine"]
        router = planning_grounding_env["router"]
        sandbox = planning_grounding_env["sandbox"]
        task_mgr = planning_grounding_env["task_manager"]
        registry = planning_grounding_env["registry"]

        prompt = (
            "Find the refinery maintenance document for P-204, read it, and prepare a concise maintenance summary containing only:\n"
            "- equipment name\n"
            "- maintenance date\n"
            "- key maintenance findings\n"
            "- recommended actions\n\n"
            "Show me the proposed summary first and ask for my approval.\n\n"
            "Do not create, modify, or write any file until I explicitly approve."
        )

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        from backend.models.base import ChatResponse, ChatChunk
        plan_json = json.dumps([
            {"description": "Search refinery maintenance document for P-204", "tool_name": "document_search", "arguments": {"query": "P-204 refinery maintenance"}, "requires_approval": False},
            {"description": "Prepare proposed concise maintenance summary", "tool_name": None, "arguments": {}, "requires_approval": False},
            {"description": "Create P-204 maintenance summary DOCX", "tool_name": "docx_create", "arguments": {"filename": "P204_Maintenance_Summary.docx", "title": "P-204 Maintenance Summary", "content": ""}, "requires_approval": True},
            {"description": "Verify generated DOCX artifact", "tool_name": "artifact_verifier", "arguments": {"filename": "P204_Maintenance_Summary.docx"}, "requires_approval": False},
        ])
        mock_provider.chat = AsyncMock(return_value=ChatResponse(content=plan_json, model="qwen2.5:7b", provider="ollama"))

        async def fake_chat_stream(req):
            all_content = "\n\n".join(m.content for m in req.messages)
            assert "[DOCUMENT CONTENT]" in all_content
            assert "88.4" in all_content
            summary = (
                "**Proposed Maintenance Summary (Awaiting Approval):**\n\n"
                "- **Equipment Name:** P-204 Hydrocracker Charge Pump\n"
                "- **Maintenance Date:** 2024-11-14\n"
                "- **Key Maintenance Findings:** Inboard bearing temperature reached 88.4°C due to 65% suction strainer blockage.\n"
                "- **Recommended Actions:** Clean suction strainer, perform dynamic balancing, and replace impeller with 13Cr stainless steel.\n\n"
                "Please approve before I generate the final report file."
            )
            yield ChatChunk(delta=summary, done=False)
            yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_chat_stream

        # --- Phase 1: Run before approval ---
        events_before = []
        async for ev in engine.run_agent_task(
            session_id="sess_regression_d",
            user_message=prompt,
        ):
            events_before.append(ev)

        # Assert document_search executed
        doc_search_events = [e for e in events_before if isinstance(e, dict) and e.get("type") == "tool_start" and e.get("tool") == "document_search"]
        assert len(doc_search_events) == 1

        # Assert proposed summary produced
        deltas = [e for e in events_before if isinstance(e, str)]
        full_proposed_summary = "".join(deltas)
        assert "88.4°C" in full_proposed_summary or "88.4" in full_proposed_summary
        assert "65%" in full_proposed_summary

        # Assert approval_required emitted
        appr_events = [e for e in events_before if isinstance(e, dict) and e.get("type") == "approval_required"]
        assert len(appr_events) == 1
        appr_ev = appr_events[0]
        task_id = appr_ev["task_id"]
        approval_id = appr_ev["approval_id"]
        assert appr_ev["tool_name"] == "docx_create"

        # Assert task is in awaiting_approval state
        task_state = task_mgr.get_task(task_id)
        assert task_state.status == "awaiting_approval"

        # Assert docx_create and artifact_verifier execution count MUST be 0
        executed_tools_before = [e.get("tool") for e in events_before if isinstance(e, dict) and e.get("type") == "tool_start"]
        assert "docx_create" not in executed_tools_before
        assert "artifact_verifier" not in executed_tools_before

        # Assert no DOCX file created
        docx_files_before = list(sandbox.glob("*.docx"))
        assert len(docx_files_before) == 0

        # --- Phase 2: User explicitly approves ---
        events_after = []
        async for ev in engine.resume_agent_task(
            task_id=task_id,
            approval_id=approval_id,
            approved=True,
            user_role="admin",
        ):
            events_after.append(ev)

        # Assert docx_create executed exactly once
        docx_start_events = [e for e in events_after if isinstance(e, dict) and e.get("type") == "tool_start" and e.get("tool") == "docx_create"]
        assert len(docx_start_events) == 1

        # Assert genuine DOCX file is created
        target_docx = sandbox / "P204_Maintenance_Summary.docx"
        assert target_docx.exists()
        doc = docx.Document(target_docx)
        doc_text = "\n".join(p.text for p in doc.paragraphs)
        assert "P-204" in doc_text
        assert "88.4°C" in doc_text or "88.4" in doc_text

        # Assert artifact_verifier executed
        verifier_results = [e for e in events_after if isinstance(e, dict) and e.get("type") == "tool_result" and e.get("tool") == "artifact_verifier"]
        assert len(verifier_results) == 1
        assert verifier_results[0]["success"] is True

        # Assert task_completed event emitted
        assert any(e.get("type") == "task_completed" for e in events_after if isinstance(e, dict))





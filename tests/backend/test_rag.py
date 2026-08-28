"""
tests/backend/test_rag.py
--------------------------
Unit tests for the RAG pipeline.
"""

import io
import pytest
from unittest.mock import AsyncMock, MagicMock
from backend.rag.ingest import DocumentParser, TextChunker, Document


class TestDocumentParser:
    def test_parse_txt(self):
        content = b"Hello world.\n\nThis is a test document.\n\nIt has three paragraphs."
        doc = DocumentParser.parse("sample.txt", content)
        assert doc.filename == "sample.txt"
        assert doc.file_type == "txt"
        assert "Hello world" in doc.text
        assert doc.document_id.startswith("doc_")

    def test_parse_md(self):
        content = b"# Title\n\nSome markdown content.\n\n## Section\n\nMore content."
        doc = DocumentParser.parse("readme.md", content)
        assert doc.file_type == "md"
        assert "markdown content" in doc.text

    def test_parse_unsupported_extension(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            DocumentParser.parse("file.exe", b"content")

    def test_parse_empty_content(self):
        with pytest.raises(ValueError, match="empty"):
            DocumentParser.parse("empty.txt", b"")

    def test_parse_utf8_fallback(self):
        content = "H\xe9llo w\xf6rld".encode("latin-1")
        doc = DocumentParser.parse("latin.txt", content)
        assert doc.file_type == "txt"
        assert len(doc.text) > 0

    def test_document_id_is_deterministic(self):
        content = b"Same content"
        id1 = DocumentParser._make_doc_id("file.txt", content)
        id2 = DocumentParser._make_doc_id("file.txt", content)
        assert id1 == id2

    def test_document_id_differs_by_name(self):
        content = b"Same content"
        id1 = DocumentParser._make_doc_id("file_a.txt", content)
        id2 = DocumentParser._make_doc_id("file_b.txt", content)
        assert id1 != id2

    def test_parse_docx(self, tmp_path):
        from docx import Document as DocxDoc
        docx_path = tmp_path / "test.docx"
        doc = DocxDoc()
        doc.add_paragraph("First paragraph about AI.")
        doc.add_paragraph("Second paragraph about data sovereignty.")
        doc.save(str(docx_path))
        content = docx_path.read_bytes()
        result = DocumentParser.parse("test.docx", content)
        assert result.file_type == "docx"
        assert "AI" in result.text or "sovereignty" in result.text

    def test_parse_pdf(self):
        pytest.importorskip("pypdf")
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        buf = io.BytesIO()
        writer.write(buf)
        content = buf.getvalue()
        with pytest.raises(ValueError, match="No text"):
            DocumentParser.parse("blank.pdf", content)


class TestTextChunker:
    def _make_doc(self, text: str) -> Document:
        return Document(
            document_id="doc_test",
            filename="test.txt",
            file_type="txt",
            text=text,
        )

    def test_chunk_short_doc(self):
        chunker = TextChunker(chunk_size=1000, chunk_overlap=100)
        doc = self._make_doc("Short document with one paragraph.")
        chunks = chunker.chunk(doc)
        assert len(chunks) >= 1
        assert chunks[0].document_id == "doc_test"

    def test_chunk_long_doc(self):
        paragraphs = [f"Paragraph {i}. " + "Content " * 50 for i in range(20)]
        text = "\n\n".join(paragraphs)
        chunker = TextChunker(chunk_size=500, chunk_overlap=50)
        doc = self._make_doc(text)
        chunks = chunker.chunk(doc)
        assert len(chunks) > 1

    def test_chunk_ids_unique(self):
        paragraphs = [f"Para {i}. " + "x " * 100 for i in range(10)]
        doc = self._make_doc("\n\n".join(paragraphs))
        chunker = TextChunker(chunk_size=300, chunk_overlap=50)
        chunks = chunker.chunk(doc)
        ids = [c.chunk_id for c in chunks]
        assert len(ids) == len(set(ids))

    def test_chunk_metadata_preserved(self):
        doc = self._make_doc("Test paragraph.")
        doc.metadata = {"source_format": "txt", "page_count": 1}
        chunker = TextChunker()
        chunks = chunker.chunk(doc)
        assert chunks[0].filename == "test.txt"
        assert chunks[0].file_type == "txt"

    def test_chunk_empty_doc(self):
        doc = self._make_doc("   \n\n   ")
        chunker = TextChunker()
        chunks = chunker.chunk(doc)
        assert chunks == []

    def test_chunk_overlap_invalid(self):
        with pytest.raises(ValueError, match="overlap"):
            TextChunker(chunk_size=100, chunk_overlap=200)

    def test_page_hint_extraction(self):
        chunker = TextChunker()
        text = "[Page 3]\nContent on page 3."
        assert chunker._extract_page_hint(text) == 3

    def test_no_page_hint(self):
        chunker = TextChunker()
        assert chunker._extract_page_hint("No page marker here.") is None


class TestDocumentServiceSecurity:
    def _make_service(self):
        from backend.config import Settings
        from backend.rag.service import DocumentService
        settings = Settings()
        embedder = AsyncMock()
        store = MagicMock()
        retriever = MagicMock()
        return DocumentService(settings, embedder, store, retriever)

    def test_path_traversal_rejected(self):
        svc = self._make_service()
        result = svc._sanitize_filename("../../etc/passwd")
        assert "/" not in result
        assert ".." not in result
        assert result == "passwd"

    def test_windows_path_traversal_rejected(self):
        svc = self._make_service()
        result = svc._sanitize_filename("..\\..\\windows\\system32\\cmd.exe")
        assert "\\" not in result
        assert ".." not in result

    def test_null_byte_rejected(self):
        svc = self._make_service()
        with pytest.raises(ValueError, match="null"):
            svc._sanitize_filename("file\x00name.txt")

    def test_normal_filename_accepted(self):
        svc = self._make_service()
        result = svc._sanitize_filename("my_document.pdf")
        assert result == "my_document.pdf"

    def test_basename_extracted(self):
        svc = self._make_service()
        result = svc._sanitize_filename("uploads/docs/report.pdf")
        assert result == "report.pdf"

    @pytest.mark.asyncio
    async def test_invalid_extension_rejected(self):
        svc = self._make_service()
        with pytest.raises(ValueError, match="not allowed"):
            await svc.ingest_document("malware.exe", b"content")

    @pytest.mark.asyncio
    async def test_empty_file_rejected(self):
        svc = self._make_service()
        with pytest.raises(ValueError):
            await svc.ingest_document("empty.txt", b"")


class TestVectorStore:
    @pytest.fixture
    def store(self, tmp_path):
        from backend.rag.store import VectorStore
        return VectorStore(persist_dir=tmp_path / "chromadb")

    def test_store_creates(self, store):
        assert store.count() == 0

    def test_add_and_count(self, store):
        store.add_chunks(
            chunk_ids=["c1", "c2"],
            embeddings=[[0.1, 0.2, 0.3], [0.4, 0.5, 0.6]],
            texts=["Text one", "Text two"],
            metadatas=[
                {"document_id": "doc1", "filename": "a.txt", "file_type": "txt", "chunk_index": 0},
                {"document_id": "doc1", "filename": "a.txt", "file_type": "txt", "chunk_index": 1},
            ],
        )
        assert store.count() == 2

    def test_list_documents(self, store):
        store.add_chunks(
            chunk_ids=["c1"],
            embeddings=[[0.1, 0.2]],
            texts=["Hello"],
            metadatas=[{"document_id": "doc1", "filename": "a.txt", "file_type": "txt", "chunk_index": 0}],
        )
        docs = store.list_documents()
        assert len(docs) == 1
        assert docs[0]["document_id"] == "doc1"

    def test_delete_document(self, store):
        store.add_chunks(
            chunk_ids=["c1", "c2"],
            embeddings=[[0.1, 0.2], [0.3, 0.4]],
            texts=["A", "B"],
            metadatas=[
                {"document_id": "doc1", "filename": "x.txt", "file_type": "txt", "chunk_index": 0},
                {"document_id": "doc1", "filename": "x.txt", "file_type": "txt", "chunk_index": 1},
            ],
        )
        deleted = store.delete_document("doc1")
        assert deleted == 2
        assert store.count() == 0

    def test_document_exists(self, store):
        assert not store.document_exists("doc_missing")
        store.add_chunks(
            chunk_ids=["c1"],
            embeddings=[[0.1]],
            texts=["hi"],
            metadatas=[{"document_id": "doc1", "filename": "f.txt", "file_type": "txt", "chunk_index": 0}],
        )
        assert store.document_exists("doc1")

    def test_upsert_idempotent(self, store):
        meta = {"document_id": "doc1", "filename": "f.txt", "file_type": "txt", "chunk_index": 0}
        store.add_chunks(["c1"], [[0.1, 0.2]], ["Text"], [meta])
        store.add_chunks(["c1"], [[0.1, 0.2]], ["Text"], [meta])
        assert store.count() == 1


class TestRetriever:
    @pytest.fixture
    def mock_embedder(self):
        e = AsyncMock()
        e.embed = AsyncMock(return_value=[0.1, 0.2, 0.3])
        return e

    @pytest.fixture
    def mock_store_empty(self):
        s = MagicMock()
        s.count.return_value = 0
        return s

    @pytest.fixture
    def mock_store_with_results(self):
        s = MagicMock()
        s.count.return_value = 2
        s.query.return_value = {
            "ids": [["c1", "c2"]],
            "documents": [["Text chunk one", "Text chunk two"]],
            "metadatas": [[
                {"document_id": "doc1", "filename": "report.pdf", "file_type": "pdf", "chunk_index": 0},
                {"document_id": "doc1", "filename": "report.pdf", "file_type": "pdf", "chunk_index": 1, "page": 3},
            ]],
            "distances": [[0.12, 0.34]],
        }
        return s

    @pytest.mark.asyncio
    async def test_empty_collection_returns_empty(self, mock_embedder, mock_store_empty):
        from backend.rag.retriever import Retriever
        r = Retriever(mock_embedder, mock_store_empty)
        results = await r.retrieve("What is AI?")
        assert results == []

    @pytest.mark.asyncio
    async def test_retrieves_chunks(self, mock_embedder, mock_store_with_results):
        from backend.rag.retriever import Retriever
        r = Retriever(mock_embedder, mock_store_with_results, top_k=5)
        results = await r.retrieve("AI question")
        assert len(results) == 2
        assert results[0].filename == "report.pdf"
        assert results[1].page == 3

    @pytest.mark.asyncio
    async def test_embed_failure_returns_empty(self, mock_store_with_results):
        from backend.rag.retriever import Retriever
        bad_embedder = AsyncMock()
        bad_embedder.embed = AsyncMock(side_effect=RuntimeError("Ollama down"))
        r = Retriever(bad_embedder, mock_store_with_results)
        results = await r.retrieve("query")
        assert results == []


class TestEngineWithRAG:
    @pytest.fixture
    def engine_with_rag(self):
        from backend.config import Settings
        from backend.models.router import ModelRouter
        from backend.agent.memory import ConversationMemory
        from backend.agent.engine import AgentEngine

        settings = Settings()
        router = MagicMock(spec=ModelRouter)
        router.default_model_id = "ollama/qwen2.5:7b"

        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        mock_response = MagicMock()
        mock_response.content = "The answer based on the document."
        mock_provider.chat = AsyncMock(return_value=mock_response)

        async def fake_stream(_req):
            from backend.models.base import ChatChunk
            yield ChatChunk(delta="The answer.", done=False)
            yield ChatChunk(delta="", done=True)

        mock_provider.chat_stream = fake_stream
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        memory = ConversationMemory()
        engine = AgentEngine(settings=settings, router=router, memory=memory)

        doc_service = MagicMock()
        doc_service.has_documents.return_value = True

        from backend.rag.retriever import RetrievedChunk
        fake_chunk = RetrievedChunk(
            text="The refund policy is 30 days.",
            document_id="doc1",
            filename="policy.pdf",
            chunk_id="doc1_chunk_0",
            chunk_index=0,
            page=1,
            score=0.15,
        )
        doc_service.retrieve = AsyncMock(return_value=[fake_chunk])
        engine.set_doc_service(doc_service)
        return engine

    @pytest.mark.asyncio
    async def test_chat_with_rag_returns_sources(self, engine_with_rag):
        content, sources = await engine_with_rag.chat("session1", "What is the refund policy?")
        assert len(sources) == 1
        assert sources[0].filename == "policy.pdf"

    @pytest.mark.asyncio
    async def test_chat_without_documents_works(self):
        from backend.config import Settings
        from backend.agent.memory import ConversationMemory
        from backend.agent.engine import AgentEngine

        settings = Settings()
        router = MagicMock()
        router.default_model_id = "ollama/qwen2.5:7b"
        mock_provider = MagicMock()
        mock_provider.provider_name = "ollama"
        mock_resp = MagicMock()
        mock_resp.content = "Hello!"
        mock_provider.chat = AsyncMock(return_value=mock_resp)
        router.get_provider_for_model.return_value = (mock_provider, "qwen2.5:7b")

        memory = ConversationMemory()
        engine = AgentEngine(settings=settings, router=router, memory=memory)

        content, sources = await engine.chat("s1", "Hello?")
        assert content == "Hello!"
        assert sources == []

    @pytest.mark.asyncio
    async def test_rag_context_not_stored_in_memory(self, engine_with_rag):
        engine = engine_with_rag
        await engine.chat("s_mem", "What is the policy?")
        history = engine._memory.get_history("s_mem")
        roles = [m.role for m in history]
        assert roles.count("system") <= 1
        assert "user" in roles
        assert "assistant" in roles
